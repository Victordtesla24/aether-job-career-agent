"""U2b (R-F4) — format-preserving tailoring engine, TDD-first (2026-08-14).

Pins the U-PLAN.md format rulings that U2b must satisfy. Written BEFORE the
implementation exists — every test in this file is expected RED against
current code, for an honest reason (a missing module/endpoint branch/field),
never a typo or a tautology. See U-PLAN.md:
  R-F4: "general format-preserving tailoring engine, NOT hand-tuned: DOCX =
  native run-level text replacement (full fidelity, flagship path); PDF =
  generalize the existing PyMuPDF block-splice with per-document geometry
  measurement + a computed fidelity confidence — low confidence ⇒ faithful
  re-render + EXPLICIT fidelity report (never silent claims); TXT/MD trivially
  preserved. Format-preservation test class: style/geometry assertions +
  zero-content-drift checks."
  R-F2: "the Format Integrity strip must state the truth for re-flowed
  documents ('rendered in Aether template; original layout preservation
  coming for this upload type')."

CONTRACT THIS FILE PINS (not yet built — defines the interface for U2b):
  * ``app.services.resume_docx.render_tailored_docx(original: bytes,
    changes: list[tuple[str, str]]) -> bytes`` — opens a real .docx, replaces
    ONLY the run text of paragraphs matching a changed bullet's original
    text, leaves every other paragraph (text + style + run formatting)
    byte-identical, and never mutates the ``original`` buffer it is given.
  * ``GET /resumes/{resume_id}/download`` grows a DOCX-native branch: when
    the resolved original (parent-then-self) is a stored DOCX (not one of
    the two bundled seed PDFs), the response is a genuine
    ``.docx`` (wordprocessingml) — never the generic re-flowed branded PDF a
    DOCX baseline gets today.
  * ``GET /resumes`` (``_with_format_preserved``) grows an HONEST
    ``formatFidelity`` object alongside the existing boolean
    ``formatPreserved`` — ``{"method": ..., "confidence": ...,
    "note": ...}`` — so a low-confidence/re-flowed résumé states the truth
    instead of a bare ``false``.
  * ``app.services.evidence_corpus.corpus_items_to_evidence_text(items) ->
    str`` — turns U2c-0's ``corpus.json`` item shape
    (``{"claim": ..., "source": ..., "confidence": ..., ...}``) into the
    plain evidence text ``ResumeTailorService.tailor``'s existing
    ``evidence_extra`` parameter already accepts, so the corpus genuinely
    widens (never bypasses) the anti-fabrication guard.

Evidence for this run: uat/reports/evidence/agents-uplift/u2b/RED-EVIDENCE.txt
"""
from __future__ import annotations

from io import BytesIO

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# A distinctive baseline body, deliberately NOT the bundled operator résumé
# (mirrors test_resume_upload.py's convention) so a round-trip can never be
# mistaken for a fallback/fixture document.
BASELINE_PARAGRAPHS = [
    ("Title", "JORDAN AKINYEMI"),
    ("Subtitle", "Senior Platform Engineer — Perth, WA, Australia"),
    ("Heading 1", "EXPERIENCE"),
    ("List Bullet", "Operated a Kubernetes fleet serving 12 million requests per day."),
    ("List Bullet", "Cut deploy lead time by 47 percent through pipeline automation."),
    ("List Bullet", "Mentored six engineers across two on-call rotations."),
]

ORIGINAL_BULLET = "Cut deploy lead time by 47 percent through pipeline automation."
TAILORED_BULLET = (
    "Cut deploy lead time by 47 percent through Docker-based pipeline automation."
)


def _make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """A real, styled .docx (python-docx) — genuine OOXML, distinct run
    formatting per paragraph so a "structure preserved" assertion is
    meaningful (a plain-text-only fixture could pass by accident)."""
    from docx import Document

    doc = Document()
    for style, text in paragraphs:
        p = doc.add_paragraph(style=style if style not in ("Title",) else None)
        run = p.add_run(text)
        if style == "Title":
            run.bold = True
            run.font.size_pt if hasattr(run.font, "size_pt") else None
        if style == "Subtitle":
            run.italic = True
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_docx(client, auth_headers, filename="jordan_resume.docx"):
    data = _make_docx_bytes(BASELINE_PARAGRAPHS)
    res = client.post(
        "/resumes/upload",
        files={"file": (filename, data, DOCX_CONTENT_TYPE)},
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert res.status_code == 201, res.text
    return res.json(), data


# ---------------------------------------------------------------------------
# (1) DOCX baseline -> tailored DOCX preserves structure (R-F4 flagship path)
# ---------------------------------------------------------------------------


def test_docx_native_render_preserves_structure_and_only_changes_target_bullet():
    """``render_tailored_docx`` (new module) must exist and, given a real
    .docx + one changed bullet, return a document with the SAME paragraph
    count/order/styles/run-formatting everywhere except the changed bullet's
    own text — the native run-level replacement R-F4 mandates.

    Expected RED today: ``app.services.resume_docx`` does not exist
    (ModuleNotFoundError) — there is no format-preserving DOCX engine yet;
    every real upload is re-flowed into the generic branded PDF template
    (routers/resumes.py's ``_branded_content`` fallback).
    """
    from docx import Document  # existing dependency; not what's under test

    from app.services.resume_docx import render_tailored_docx  # NEW module (R-F4)

    original_bytes = _make_docx_bytes(BASELINE_PARAGRAPHS)
    original_snapshot = bytes(original_bytes)  # independent copy for the immutability check

    output = render_tailored_docx(original_bytes, [(ORIGINAL_BULLET, TAILORED_BULLET)])

    assert isinstance(output, (bytes, bytearray))
    assert output != original_bytes, "a genuine tailoring edit must change the output bytes"
    assert original_bytes == original_snapshot, (
        "render_tailored_docx must never mutate the ORIGINAL bytes buffer it "
        "was given — the baseline is immutable even at the in-memory level"
    )

    result_doc = Document(BytesIO(output))
    source_doc = Document(BytesIO(original_snapshot))
    result_paragraphs = result_doc.paragraphs
    source_paragraphs = source_doc.paragraphs

    assert len(result_paragraphs) == len(source_paragraphs), (
        "paragraph count/order must be preserved — R-F4 forbids re-flowing "
        "into a different structure"
    )

    changed_count = 0
    for src_p, out_p in zip(source_paragraphs, result_paragraphs):
        assert out_p.style.name == src_p.style.name, (
            f"paragraph style must be preserved unchanged "
            f"(source={src_p.style.name!r}, output={out_p.style.name!r})"
        )
        if src_p.text.strip() == ORIGINAL_BULLET:
            assert out_p.text.strip() == TAILORED_BULLET, (
                "the targeted bullet's TEXT must be updated to the tailored wording"
            )
            changed_count += 1
        else:
            assert out_p.text == src_p.text, (
                f"an UNCHANGED paragraph's text must stay byte-identical "
                f"(got {out_p.text!r} vs {src_p.text!r})"
            )
            # Run-level formatting on untouched paragraphs must be untouched too.
            assert len(out_p.runs) == len(src_p.runs)
            for src_run, out_run in zip(src_p.runs, out_p.runs):
                assert out_run.bold == src_run.bold
                assert out_run.italic == src_run.italic

    assert changed_count == 1, "exactly one paragraph (the changed bullet) may differ"


# ---------------------------------------------------------------------------
# (4) ATS-safety constraints hold on the DOCX-native output
# ---------------------------------------------------------------------------


def test_docx_native_output_stays_ats_parseable_and_scores_the_new_keyword():
    """The tailored DOCX's text layer must remain real, extractable text (no
    rasterised/embedded-image substitute that would hide it from an ATS
    parser), and the JD keyword the rewrite added must be genuinely
    extractable and score-visible via the SAME ATS engine the product uses
    everywhere else.

    Expected RED today: ``app.services.resume_docx`` does not exist.
    """
    from docx import Document

    from app.services.ats_engine import ATSEngine
    from app.services.resume_docx import render_tailored_docx  # NEW module (R-F4)

    original_bytes = _make_docx_bytes(BASELINE_PARAGRAPHS)
    output = render_tailored_docx(original_bytes, [(ORIGINAL_BULLET, TAILORED_BULLET)])

    doc = Document(BytesIO(output))
    assert len(doc.inline_shapes) == 0, (
        "no image/OLE object may be introduced — ATS parsers read text runs, "
        "not pictures"
    )
    extracted = "\n".join(p.text for p in doc.paragraphs)
    assert "Docker-based pipeline automation" in extracted, (
        "the tailored keyword must land in a genuinely extractable text run"
    )

    job_description = "Looking for an engineer experienced with Docker and Kubernetes."
    before_score = ATSEngine().score(
        "\n".join(text for _, text in BASELINE_PARAGRAPHS), job_description
    )
    after_score = ATSEngine().score(extracted, job_description)
    assert after_score.keyword_match >= before_score.keyword_match, (
        "the format-preserving DOCX render must not regress the real ATS "
        "keyword score relative to the baseline text"
    )


# ---------------------------------------------------------------------------
# (3) Baseline bytes remain untouched by tailoring / immutability regression,
#     exercised end-to-end through the /download endpoint's new DOCX branch.
# ---------------------------------------------------------------------------


def test_download_of_docx_tailored_version_returns_native_docx_and_baseline_stays_pristine(
    client, auth_headers,
):
    """End-to-end: uploading a DOCX baseline, creating a tailored child
    version, and downloading it must (a) return a genuine, format-preserved
    ``.docx`` — not the generic re-flowed branded PDF every non-bundled
    upload gets today — and (b) leave the baseline's own stored bytes
    byte-identical, confirmed via ``GET /resumes/{id}/original``.

    Expected RED today: ``/download`` has no DOCX-native branch — a
    DOCX-sourced tailored version has no bundled source PDF match, so it
    falls to ``_branded_content`` + ``create_branded_resume_pdf`` and comes
    back as ``application/pdf`` (format lost), not a ``.docx``.
    """
    baseline, original_bytes = _upload_docx(client, auth_headers)
    baseline_id = baseline["id"]

    child_sections = dict(baseline["sections"])
    child_sections["bullets"] = [
        {
            "text": TAILORED_BULLET if b["text"].strip() == ORIGINAL_BULLET else b["text"],
            "evidenceRef": b["evidenceRef"],
        }
        for b in baseline["sections"]["bullets"]
    ]
    from app.repositories.resume import ResumeRepository as _Repo

    child = _Repo().create(
        _test_user_id(client, auth_headers),
        child_sections,
        baseline["formatHash"],  # unchanged hash: no new original bytes for a child
        label="Tailored — Platform Engineer @ Acme",
        version=_Repo().next_version(_test_user_id(client, auth_headers)),
        parent_id=baseline_id,
    )

    res = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert res.status_code == 200, res.text
    assert res.headers["content-type"].startswith(DOCX_CONTENT_TYPE), (
        f"a DOCX-sourced tailored version must download as a native .docx, "
        f"got content-type={res.headers['content-type']!r} — the format was "
        f"lost to the generic re-flowed PDF fallback"
    )
    from docx import Document

    result_doc = Document(BytesIO(res.content))
    assert any(
        TAILORED_BULLET in p.text for p in result_doc.paragraphs
    ), "the downloaded native docx must contain the tailored bullet text"

    original_after = client.get(f"/resumes/{baseline_id}/original", headers=auth_headers)
    assert original_after.status_code == 200, original_after.text
    assert original_after.content == original_bytes, (
        "downloading a tailored DOCX version must never mutate the baseline's "
        "own immutable stored bytes"
    )


def _test_user_id(client, auth_headers) -> str:
    me = client.get("/auth/me", headers=auth_headers)
    assert me.status_code == 200, me.text
    return me.json()["id"]


# ---------------------------------------------------------------------------
# (5) Honest degradation: a real (non-bundled) PDF baseline must carry an
#     EXPLICIT fidelity report, never a silent re-format (R-F2/R-F4).
# ---------------------------------------------------------------------------


def _make_pdf_bytes(lines: list[str]) -> bytes:
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(612, 792))
    y = 740
    for line in lines:
        c.drawString(72, y, line)
        y -= 16
    c.showPage()
    c.save()
    return buf.getvalue()


def test_reflowed_pdf_baseline_carries_an_explicit_honest_fidelity_report(
    client, auth_headers,
):
    """A real user-uploaded PDF (not one of the two bundled seed PDFs) is
    re-flowed on download today, and ``formatPreserved`` already says
    ``False`` (MON-011) — but R-F2/R-F4 require an EXPLICIT report of *why*
    and *what will actually happen*, not just a bare boolean. ``GET
    /resumes`` must stamp a ``formatFidelity`` object with a low/non-high
    confidence and an honest note naming the re-flow — never silence.

    Expected RED today: ``_with_format_preserved`` stamps only the boolean
    ``formatPreserved``; there is no ``formatFidelity`` key at all.
    """
    pdf_bytes = _make_pdf_bytes([
        "RILEY NAKAMURA",
        "Staff Data Scientist — Adelaide, SA, Australia",
        "",
        "EXPERIENCE",
        "- Built a fraud model serving 40 million scored transactions weekly.",
    ])
    res = client.post(
        "/resumes/upload",
        files={"file": ("riley_resume.pdf", pdf_bytes, "application/pdf")},
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert res.status_code == 201, res.text
    resume_id = res.json()["id"]

    listing = client.get("/resumes", headers=auth_headers).json()
    stamped = next(r for r in listing if r["id"] == resume_id)

    assert stamped["formatPreserved"] is False  # MON-011 baseline behaviour, still true

    assert "formatFidelity" in stamped, (
        "R-F2/R-F4: a re-flowed résumé must carry an EXPLICIT fidelity "
        "report object, not just a bare boolean"
    )
    fidelity = stamped["formatFidelity"]
    assert fidelity["confidence"] in ("low", "medium"), (
        f"a re-flowed (non-native) render is not high-confidence fidelity, "
        f"got {fidelity!r}"
    )
    note = str(fidelity.get("note", "")).lower()
    assert "template" in note or "re-flow" in note or "reflow" in note, (
        f"the note must honestly name what actually happens to this upload "
        f"type, got {fidelity!r}"
    )
    # Never a silent/fabricated preservation claim.
    assert "byte-for-byte" not in note and "pixel-perfect" not in note


# ---------------------------------------------------------------------------
# (2) Tailored content claims trace to corpus/baseline evidence — a claim the
#     U2c-0 corpus does NOT support must be rejected/flagged, never included.
# ---------------------------------------------------------------------------


_MOCK_CORPUS_ITEMS = [
    {
        "id": "c1",
        "source": "github:aether-job-career-agent",
        "category": "skill",
        "claim": "Built a multi-agent orchestration pipeline using Terraform for IaC.",
        "stated_or_inferred": "inferred",
        "confidence": "high",
        "sourceUrl": "https://github.com/example/aether-job-career-agent",
        "note": "byte-share, not a direct authorship claim",
    },
    {
        "id": "c2",
        "source": "portfolio",
        "category": "skill",
        "claim": "Skills & Certifications lists Docker and Kubernetes.",
        "stated_or_inferred": "stated",
        "confidence": "high",
        "sourceUrl": "https://forgotten-mistory.web.app/",
        "note": "",
    },
]


def test_corpus_adapter_turns_u2c0_items_into_evidence_text_supporting_a_tailor_claim():
    """``corpus_items_to_evidence_text`` (new adapter, U2c-0 -> tailoring
    glue) must exist and turn corpus items into the plain evidence text
    ``ResumeTailorService`` already accepts as ``evidence_extra`` — so a
    rewrite that surfaces a corpus-only claim (Terraform, present ONLY in the
    mock corpus, never in the résumé text) is ACCEPTED, not reverted.

    Expected RED today: ``app.services.evidence_corpus`` does not exist —
    there is no glue between the U2c-0 corpus.json shape and the tailoring
    guard's evidence_extra parameter.
    """
    from app.services.evidence_corpus import corpus_items_to_evidence_text  # NEW

    evidence_text = corpus_items_to_evidence_text(_MOCK_CORPUS_ITEMS)
    assert "Terraform" in evidence_text
    assert "Docker" in evidence_text

    from app.services.resume_tailor import ResumeTailorService

    resume_text = (
        "JORDAN AKINYEMI\nSenior Platform Engineer\n\nEXPERIENCE\n"
        "- Operated a Kubernetes fleet serving 12 million requests per day.\n"
    )
    originals = [
        {"text": "Operated a Kubernetes fleet serving 12 million requests per day.",
         "evidenceRef": "bullet-0"},
    ]
    raw = {
        "bullets": [
            {
                "text": (
                    "Operated a Kubernetes fleet serving 12 million requests per "
                    "day, provisioned via Terraform."
                ),
                "evidenceRef": "bullet-0",
            }
        ],
        "evidenceRefs": ["bullet-0"],
    }

    service = ResumeTailorService.__new__(ResumeTailorService)  # no LLM client needed
    result = service._validate(
        raw, originals, resume_text, job_description="", evidence_extra=evidence_text,
    )
    assert result.rejected == [], (
        f"a claim genuinely supported by the U2c-0 corpus must be ACCEPTED, "
        f"got rejected={result.rejected!r}"
    )
    assert result.bullets[0]["text"].strip().endswith("Terraform.")


def test_corpus_adapter_still_rejects_a_claim_the_corpus_does_not_support():
    """The mirror case: a claim absent from BOTH the résumé and the mock
    corpus (a fabricated skill, e.g. 'Rust') must still be rejected/reverted
    — the corpus EXTENDS the evidence base, it never bypasses the guard.

    Expected RED today: same missing ``app.services.evidence_corpus`` module.
    """
    from app.services.evidence_corpus import corpus_items_to_evidence_text  # NEW

    evidence_text = corpus_items_to_evidence_text(_MOCK_CORPUS_ITEMS)
    assert "Rust" not in evidence_text

    from app.services.resume_tailor import ResumeTailorService

    resume_text = (
        "JORDAN AKINYEMI\nSenior Platform Engineer\n\nEXPERIENCE\n"
        "- Operated a Kubernetes fleet serving 12 million requests per day.\n"
    )
    originals = [
        {"text": "Operated a Kubernetes fleet serving 12 million requests per day.",
         "evidenceRef": "bullet-0"},
    ]
    raw = {
        "bullets": [
            {
                "text": (
                    "Operated a Kubernetes fleet serving 12 million requests per "
                    "day, using a custom Rust control plane."
                ),
                "evidenceRef": "bullet-0",
            }
        ],
        "evidenceRefs": ["bullet-0"],
    }

    service = ResumeTailorService.__new__(ResumeTailorService)
    result = service._validate(
        raw, originals, resume_text, job_description="", evidence_extra=evidence_text,
    )
    assert result.rejected != [], (
        "a fabricated claim unsupported by resume OR corpus must be rejected"
    )
    assert result.bullets[0]["text"].strip() == originals[0]["text"], (
        "an unsupported rewrite must revert to the ORIGINAL bullet, never "
        "ship silently"
    )
