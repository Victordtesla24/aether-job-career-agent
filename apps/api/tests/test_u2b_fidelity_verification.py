"""U2b truth round — format fidelity must be VERIFIED, never asserted.

Live production evidence (uat/reports/evidence/agents-uplift/u2b/verify/,
2026-08-14) falsified the shipped claim: tailored résumé
``c34ec9016096f3ad0ec06a733`` reported ``formatFidelity = {"method":
"pdf-in-place-splice", "confidence": "high", "note": "Only the reworded
bullets are redrawn on your original PDF — every other element is identical to
the source document."}`` while the file the user actually downloads still
contained the ORIGINAL text of one of the four reworded bullets. The in-place
splice engine only edits right-column work bullets (``resume_pdf.py``
``_RIGHT_COL_MIN_X = 225.0``); the rewritten bullet lives in the left rail at
x0 ≈ 46.5, so it was silently skipped — and the fidelity report, which is
computed from the resume's *metadata* (hash match + content type) rather than
from the produced document, claimed completeness anyway.

The contract these tests pin:

1. Every fidelity claim about a tailored download is derived by RE-READING the
   artifact that was just produced and checking each requested change is
   really in it (``verification: "post-render-text-extraction"``), with the
   applied/dropped counts exposed in the API payload.
2. A change the engine could not apply is NAMED, not hidden, and drops the
   confidence off "high".
3. The listing (``GET /resumes``) never asserts completeness it has not
   verified for a tailored version.
4. The DOCX-native path's report is verification-derived too — not the
   renderer's own bookkeeping of what it *believes* it replaced.

Every fixture here is derived from the repository's own bundled résumé asset
at run time (never hard-coded prose), so the tests exercise the real geometry
that produced the live defect.
"""
from __future__ import annotations

import hashlib
from io import BytesIO

import pytest

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _user_id(client, auth_headers) -> str:
    me = client.get("/auth/me", headers=auth_headers)
    assert me.status_code == 200, me.text
    return me.json()["id"]


def _bundled_pdf():
    from app.agents.fit_scorer import get_base_resume_path

    return get_base_resume_path()


def _left_rail_line() -> str:
    """A real line the in-place splice engine structurally cannot edit.

    Read from the bundled résumé itself: any text line left of
    ``resume_pdf._RIGHT_COL_MIN_X`` is in the two-column layout's left rail,
    which ``_detect_blocks`` deliberately excludes. Longest such line wins so
    the fixture is deterministic and unambiguous.
    """
    import fitz

    from app.services.resume_pdf import _RIGHT_COL_MIN_X

    doc = fitz.open(_bundled_pdf())
    try:
        candidates: list[str] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    spans = line["spans"]
                    if not spans:
                        continue
                    if min(s["bbox"][0] for s in spans) >= _RIGHT_COL_MIN_X:
                        continue
                    text = " ".join("".join(s["text"] for s in spans).split())
                    if len(text) > 25:
                        candidates.append(text)
    finally:
        doc.close()
    assert candidates, "the bundled résumé must have left-rail text to fixture on"
    return max(candidates, key=len)


def _right_column_bullet() -> str:
    """A real work bullet the splice engine CAN edit (right column)."""
    from app.services.resume_pdf import extract_pdf_bullets

    bullets = [b for b in extract_pdf_bullets(_bundled_pdf()) if ":" in b and len(b) > 80]
    assert bullets, "the bundled résumé must have right-column work bullets"
    return bullets[0]


def _seed_bundled_baseline(client, auth_headers) -> tuple[dict, str, str]:
    """A baseline whose formatHash matches the bundled PDF, plus its two bullets.

    Returns ``(baseline_row, left_rail_bullet, right_column_bullet)``.
    """
    from app.repositories.resume import ResumeRepository

    left = _left_rail_line()
    right = _right_column_bullet()
    user_id = _user_id(client, auth_headers)
    repo = ResumeRepository()
    baseline = repo.create(
        user_id,
        {
            "raw_text": f"{left}\n{right}",
            "bullets": [
                {"text": left, "evidenceRef": "bullet-0"},
                {"text": right, "evidenceRef": "bullet-1"},
            ],
            "contact": {},
        },
        hashlib.sha256(_bundled_pdf().read_bytes()).hexdigest(),
        label="Baseline — bundled layout",
        version=repo.next_version(user_id),
    )
    return baseline, left, right


def _tailor_child(client, auth_headers, baseline: dict, rewrites: dict[str, str]) -> dict:
    """A tailored child version whose bullets carry ``rewrites`` by evidenceRef."""
    from app.repositories.resume import ResumeRepository

    user_id = _user_id(client, auth_headers)
    repo = ResumeRepository()
    bullets = [
        {
            "text": rewrites.get(b["evidenceRef"], b["text"]),
            "evidenceRef": b["evidenceRef"],
        }
        for b in baseline["sections"]["bullets"]
    ]
    sections = dict(baseline["sections"])
    sections["bullets"] = bullets
    return repo.create(
        user_id,
        sections,
        baseline["formatHash"],
        label="Tailored — Delivery Manager @ Nearmap",
        version=repo.next_version(user_id),
        parent_id=baseline["id"],
    )


def _pdf_text(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        return " ".join(" ".join(page.get_text().split()) for page in doc)
    finally:
        doc.close()


def _make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    from docx import Document

    doc = Document()
    for style, text in paragraphs:
        paragraph = doc.add_paragraph(style=style or None)
        paragraph.add_run(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# (1) The live defect: a change the splice cannot apply must be REPORTED
# ---------------------------------------------------------------------------


def test_pdf_splice_fidelity_names_the_change_it_could_not_apply(client, auth_headers):
    """The exact production failure, reproduced against the same bundled layout.

    One rewrite targets a left-rail line the splice engine skips; the other a
    right-column work bullet it can redraw. The fidelity report must be
    computed from the produced PDF — 1 of 2 applied, the dropped one named —
    instead of claiming high-confidence completeness for both.
    """
    baseline, left, right = _seed_bundled_baseline(client, auth_headers)
    left_after = f"Product and delivery leadership across regulated platforms; {left}"
    head, _, tail = right.partition(":")
    right_after = f"{head}: Re-scoped for the target role, {tail.strip()}"
    child = _tailor_child(
        client, auth_headers, baseline, {"bullet-0": left_after, "bullet-1": right_after}
    )

    res = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers)
    assert res.status_code == 200, res.text
    report = res.json()

    assert report["method"] == "pdf-in-place-splice"
    assert report["verification"] == "post-render-text-extraction", (
        "the claim must come from re-reading the produced document, not from "
        f"resume metadata, got {report!r}"
    )
    assert report["changesRequested"] == 2
    assert report["changesApplied"] == 1, report
    assert report["changesDropped"] == 1, report
    assert report["confidence"] != "high", (
        "a download that silently drops a tailored change is not a "
        f"high-confidence fidelity claim, got {report!r}"
    )
    note = report["note"]
    assert "1 of 2" in note, f"the note must state the honest counts, got {note!r}"
    assert "could not be applied" in note.lower(), note
    dropped = report["droppedChanges"]
    assert len(dropped) == 1 and dropped[0]["after"].startswith(left_after[:40]), dropped

    # The report must match the file the user really gets — verified here
    # independently of the endpoint's own machinery.
    download = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert download.status_code == 200, download.text
    text = _pdf_text(download.content)
    assert left_after[:60] not in text, (
        "fixture invalid: the left-rail rewrite must genuinely be missing"
    )
    assert right_after.partition(":")[2].strip()[:60] in text, (
        "fixture invalid: the right-column rewrite must genuinely be applied"
    )
    assert download.headers["X-Aether-Changes-Applied"] == "1"
    assert download.headers["X-Aether-Changes-Dropped"] == "1"


def test_listing_does_not_claim_completeness_it_has_not_verified(client, auth_headers):
    """``GET /resumes`` describes the mechanism; it never asserts the outcome.

    The listing cannot re-render every version, so for a tailored row it must
    say the per-change verification is pending — not repeat the unconditional
    "every other element is identical" claim the live run falsified.
    """
    baseline, left, right = _seed_bundled_baseline(client, auth_headers)
    child = _tailor_child(
        client, auth_headers, baseline, {"bullet-0": f"Rewritten: {left}"}
    )

    listing = client.get("/resumes", headers=auth_headers).json()
    row = next(r for r in listing if r["id"] == child["id"])
    fidelity = row["formatFidelity"]

    assert fidelity["method"] == "pdf-in-place-splice"
    assert fidelity["confidence"] == "unverified", (
        f"an unverified tailored row must say so, got {fidelity!r}"
    )
    assert "every other element is identical" not in fidelity["note"].lower(), (
        f"the falsified completeness claim must not survive, got {fidelity!r}"
    )
    assert "verif" in fidelity["note"].lower(), fidelity


# ---------------------------------------------------------------------------
# (3) The DOCX-native path is verification-derived too
# ---------------------------------------------------------------------------


def test_docx_native_fidelity_is_derived_from_the_produced_document(
    client, auth_headers,
):
    original = "Led a payments migration serving 4 million monthly customers."
    rewritten = "Led a payments platform migration serving 4 million monthly customers."
    upload = client.post(
        "/resumes/upload",
        files={
            "file": (
                "priya_resume.docx",
                _make_docx_bytes([
                    ("", "PRIYA RAGHAVAN"),
                    ("Heading 1", "EXPERIENCE"),
                    ("List Bullet", original),
                    ("List Bullet", "Cut incident response time by 35 percent."),
                ]),
                DOCX_CONTENT_TYPE,
            )
        },
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert upload.status_code == 201, upload.text
    baseline = upload.json()
    ref = next(
        b["evidenceRef"]
        for b in baseline["sections"]["bullets"]
        if b["text"].strip().endswith(original)
    )
    child = _tailor_child(client, auth_headers, baseline, {ref: rewritten})

    res = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers)
    assert res.status_code == 200, res.text
    report = res.json()

    assert report["method"] == "docx-native"
    assert report["verification"] == "post-render-text-extraction", report
    assert report["changesRequested"] == 1
    assert report["changesApplied"] == 1, report
    assert report["changesDropped"] == 0, report
    assert report["confidence"] == "high", report


def test_fidelity_endpoint_is_owner_only(client, auth_headers):
    baseline, _left, _right = _seed_bundled_baseline(client, auth_headers)
    other = client.post(
        "/auth/register",
        json={"email": "u2b-fidelity-other@example.com", "password": "Sup3rSecret"},
    )
    assert other.status_code in (200, 201, 409), other.text
    login = client.post(
        "/auth/login",
        json={"email": "u2b-fidelity-other@example.com", "password": "Sup3rSecret"},
    )
    assert login.status_code == 200, login.text
    headers = {"Authorization": f"Bearer {login.json()['access_token']}"}

    res = client.get(f"/resumes/{baseline['id']}/fidelity", headers=headers)
    assert res.status_code == 404, res.text


# ---------------------------------------------------------------------------
# Verifier unit contract: a bold lead-in must not read as a dropped change
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("media_type", "unreadable"),
    [
        ("application/pdf", b"not a real pdf"),
        (DOCX_CONTENT_TYPE, b"not a real docx package"),
        ("text/plain; charset=utf-8", b"\xff\xfe\x00not utf-8"),
    ],
)
def test_verifier_reports_unverified_for_an_unreadable_artifact(media_type, unreadable):
    """An artifact that cannot be re-read is UNVERIFIED, never "all dropped".

    Counting every change as dropped would be its own fabrication — the
    opposite lie to the one this round fixes.
    """
    from app.services.format_verification import verify_changes

    result = verify_changes(unreadable, media_type, [("before", "after")])
    assert result.text_extracted is False
    assert result.requested == 1
    assert result.applied_count == 0
    assert result.dropped_count == 0
    assert result.complete is False


def test_verifier_counts_a_bold_lead_in_rewrite_as_applied():
    """A rewrite split across two PDF text writers is APPLIED, not dropped.

    ``resume_pdf._render_block`` draws a bullet's bold lead-in and its grey
    body through separate ``TextWriter``s, so the extracted text stream can
    interleave other page content between them. Exact-substring matching
    reported the real production artifact's applied bold-lead-in rewrite as
    missing (coverage 0.917) — this pins the coverage rule that fixes it.
    """
    from app.services.format_verification import verify_changes
    from app.services.resume_pdf import extract_pdf_bullets, render_tailored_pdf

    bullet = _right_column_bullet()
    head, _, tail = bullet.partition(":")
    after = f"{head}: Re-scoped for the target role, {tail.strip()}"
    rendered = render_tailored_pdf(_bundled_pdf(), [(bullet, after)])

    result = verify_changes(rendered, "application/pdf", [(bullet, after)])
    assert result.text_extracted is True
    assert result.applied_count == 1, [o.coverage for o in result.outcomes]
    assert extract_pdf_bullets(_bundled_pdf()), "fixture sanity: bundled bullets exist"
