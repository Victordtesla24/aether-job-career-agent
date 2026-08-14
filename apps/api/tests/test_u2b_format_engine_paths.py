"""U2b — the format-preserving paths the R-F4 contract tests do not reach.

``test_u2b_format_engine.py`` pins the DOCX flagship path, the honest fidelity
report for a re-flowed PDF, and the corpus→guard adapter. This file covers the
rest of the same slice, so no branch of the engine ships unexercised:

* the **plain-text/Markdown** path R-F4 calls "trivially preserved" — the
  user's own file back, with only the reworded lines changed;
* a **base .docx download**, which must be the user's stored bytes verbatim
  (nothing is re-rendered when there is nothing to rewrite);
* the **fidelity report** for a DOCX-backed version (``formatPreserved: True``
  with a ``docx-native`` method) and the honest **unknown** third state;
* the **evidence-corpus store**: idempotent import of a U2c-0 snapshot,
  JD-ranked bounded evidence text, and the honest empty state.
"""
from __future__ import annotations

import json
import uuid
from io import BytesIO

import pytest

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

MD_BASELINE = """# ALEX MERCADO
Senior Data Engineer — Hobart, TAS, Australia

## EXPERIENCE
- Rebuilt an ingestion pipeline that moved 30 terabytes of telemetry weekly.
- Reduced warehouse spend by 22 percent through partition tuning.
"""

MD_ORIGINAL_BULLET = (
    "Rebuilt an ingestion pipeline that moved 30 terabytes of telemetry weekly."
)
MD_TAILORED_BULLET = (
    "Rebuilt an Airflow ingestion pipeline that moved 30 terabytes of telemetry weekly."
)


def _user_id(client, auth_headers) -> str:
    me = client.get("/auth/me", headers=auth_headers)
    assert me.status_code == 200, me.text
    return me.json()["id"]


def _make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    from docx import Document

    doc = Document()
    for style, text in paragraphs:
        paragraph = doc.add_paragraph(style=style or None)
        paragraph.add_run(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


DOCX_PARAGRAPHS = [
    ("", "PRIYA RAGHAVAN"),
    ("Heading 1", "EXPERIENCE"),
    ("List Bullet", "Led a payments migration serving 4 million monthly customers."),
    ("List Bullet", "Cut incident response time by 35 percent with runbook automation."),
]


def _upload(client, auth_headers, filename: str, data: bytes, content_type: str):
    res = client.post(
        "/resumes/upload",
        files={"file": (filename, data, content_type)},
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert res.status_code == 201, res.text
    return res.json()


# ---------------------------------------------------------------------------
# Markdown/plain-text: trivially preserved (R-F4)
# ---------------------------------------------------------------------------


def test_markdown_baseline_downloads_natively_with_only_the_reworded_line_changed(
    client, auth_headers,
):
    baseline = _upload(
        client, auth_headers, "alex_resume.md", MD_BASELINE.encode(), "text/markdown"
    )
    bullets = baseline["sections"]["bullets"]
    assert any(b["text"].strip() == MD_ORIGINAL_BULLET for b in bullets), (
        f"the markdown bullets must be extracted for tailoring, got {bullets!r}"
    )

    from app.repositories.resume import ResumeRepository

    repo = ResumeRepository()
    user_id = _user_id(client, auth_headers)
    child_sections = dict(baseline["sections"])
    child_sections["bullets"] = [
        {
            "text": MD_TAILORED_BULLET
            if b["text"].strip() == MD_ORIGINAL_BULLET
            else b["text"],
            "evidenceRef": b["evidenceRef"],
        }
        for b in bullets
    ]
    child = repo.create(
        user_id,
        child_sections,
        baseline["formatHash"],
        label="Tailored — Data Engineer @ Acme",
        version=repo.next_version(user_id),
        parent_id=baseline["id"],
    )

    # Base: the user's own file, byte-identical.
    base_download = client.get(
        f"/resumes/{baseline['id']}/download", headers=auth_headers
    )
    assert base_download.status_code == 200, base_download.text
    assert base_download.content == MD_BASELINE.encode()

    # Tailored: same file, one sentence changed, every other line untouched.
    tailored = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert tailored.status_code == 200, tailored.text
    assert tailored.headers["content-type"].startswith("text/markdown")
    text = tailored.content.decode()
    assert MD_TAILORED_BULLET in text
    assert MD_ORIGINAL_BULLET not in text
    assert "# ALEX MERCADO" in text and "## EXPERIENCE" in text
    assert "Reduced warehouse spend by 22 percent through partition tuning." in text


# ---------------------------------------------------------------------------
# DOCX base download + fidelity reporting
# ---------------------------------------------------------------------------


def test_docx_base_download_returns_the_stored_bytes_verbatim(client, auth_headers):
    data = _make_docx_bytes(DOCX_PARAGRAPHS)
    baseline = _upload(client, auth_headers, "priya_resume.docx", data, DOCX_CONTENT_TYPE)

    res = client.get(f"/resumes/{baseline['id']}/download", headers=auth_headers)
    assert res.status_code == 200, res.text
    assert res.headers["content-type"].startswith(DOCX_CONTENT_TYPE)
    assert res.content == data, (
        "a base .docx has nothing to rewrite, so its download must be the "
        "user's own stored document, unmodified"
    )


def test_docx_upload_extracts_word_list_items_as_bullets(client, auth_headers):
    """Word stores a bullet glyph in numbering.xml, not in the paragraph text.

    Without marking list paragraphs during extraction, a .docx résumé yields
    ZERO tailorable bullets — the upload succeeds and nothing in it can ever be
    rewritten.
    """
    baseline = _upload(
        client,
        auth_headers,
        "priya_resume.docx",
        _make_docx_bytes(DOCX_PARAGRAPHS),
        DOCX_CONTENT_TYPE,
    )
    texts = [b["text"].strip() for b in baseline["sections"]["bullets"]]
    assert "Led a payments migration serving 4 million monthly customers." in texts
    assert "Cut incident response time by 35 percent with runbook automation." in texts


def test_docx_resume_reports_docx_native_fidelity_and_preserved_true(
    client, auth_headers,
):
    baseline = _upload(
        client,
        auth_headers,
        "priya_resume.docx",
        _make_docx_bytes(DOCX_PARAGRAPHS),
        DOCX_CONTENT_TYPE,
    )
    listing = client.get("/resumes", headers=auth_headers).json()
    stamped = next(r for r in listing if r["id"] == baseline["id"])

    assert stamped["formatPreserved"] is True
    assert stamped["formatFidelity"]["method"] == "docx-native"
    assert stamped["formatFidelity"]["confidence"] == "high"
    assert "native document editing" in stamped["formatFidelity"]["note"].lower()


def test_unresolvable_source_version_reports_unknown_not_a_guess():
    """The honest third state: a version whose parent cannot be read.

    ``formatPreserved`` is ``None`` — neither an affirmative claim nor a denial
    — and the fidelity note says so, because guessing either way would be a
    fabricated statement about the user's own document.
    """
    from app.services.resume_format import stamp_fidelity

    orphan = {"id": "child", "parentId": "a-version-not-in-this-listing", "formatHash": "h"}
    stamped = stamp_fidelity([orphan], set(), {})[0]

    assert stamped["formatPreserved"] is None
    assert stamped["formatFidelity"]["method"] == "unknown"
    assert stamped["formatFidelity"]["confidence"] == "unknown"
    assert "could not be resolved" in stamped["formatFidelity"]["note"]


# ---------------------------------------------------------------------------
# Evidence corpus store (U2c-0 snapshot -> app data layer)
# ---------------------------------------------------------------------------


_SNAPSHOT = [
    {
        "id": "i1",
        "claim": "Delivered a Kubernetes platform migration for 4 million customers.",
        "category": "achievement",
        "source": "baseline",
        "sourceUrl": "https://example.invalid/resume",
        "stated_or_inferred": "stated",
        "confidence": "high",
        "asOf": "2026-08-14T00:21:17Z",
        "note": "Verbatim from baseline resume.",
    },
    {
        "id": "i2",
        "claim": "Uses Terraform to provision infrastructure in a public repo.",
        "category": "skill",
        "source": "github:platform-iac",
        "sourceUrl": "https://example.invalid/repo",
        "stated_or_inferred": "inferred",
        "confidence": "med",
        "asOf": "2026-08-14T00:22:00Z",
        "note": "byte-share, not a direct authorship claim",
    },
    {"id": "", "claim": "no id — must be skipped", "source": "portfolio"},
]


@pytest.fixture()
def corpus_user() -> str:
    return f"u2b-corpus-{uuid.uuid4().hex[:12]}"


def test_corpus_import_is_idempotent_and_round_trips_provenance(tmp_path, corpus_user):
    from app.repositories.evidence_corpus import EvidenceCorpusRepository
    from app.services.evidence_corpus import import_corpus_file

    path = tmp_path / "corpus.json"
    path.write_text(json.dumps(_SNAPSHOT), encoding="utf-8")

    assert import_corpus_file(corpus_user, path) == 2, (
        "an item with no id is not addressable evidence and must be skipped"
    )
    assert import_corpus_file(corpus_user, path) == 2  # idempotent, not duplicated

    rows = EvidenceCorpusRepository().list_by_user(corpus_user)
    assert len(rows) == 2
    by_id = {row["itemId"]: row for row in rows}
    assert by_id["i2"]["statedOrInferred"] == "inferred"
    assert by_id["i2"]["confidence"] == "med"
    assert by_id["i2"]["sourceUrl"] == "https://example.invalid/repo"
    assert by_id["i2"]["source"] == "github:platform-iac"


def test_corpus_evidence_is_jd_ranked_bounded_and_url_free(tmp_path, corpus_user):
    from app.services.evidence_corpus import build_corpus_evidence, import_corpus_file

    path = tmp_path / "corpus.json"
    path.write_text(json.dumps(_SNAPSHOT), encoding="utf-8")
    import_corpus_file(corpus_user, path)

    evidence = build_corpus_evidence(
        corpus_user, "Platform engineer fluent in Terraform and cloud provisioning."
    )
    assert "Terraform" in evidence and "Kubernetes" in evidence
    # Source URLs must never enter the fabrication index — every stem in this
    # text becomes a licensed token for the guard.
    assert "https://" not in evidence
    # JD-relevant evidence ranks first.
    assert evidence.index("Terraform") < evidence.index("Kubernetes")
    # Units are blank-line separated: the shape the guard's context scoping splits on.
    assert "\n\n" in evidence

    # The character budget is a hard cap: a unit that does not fit is SKIPPED,
    # never truncated mid-claim (half a claim is not evidence). 150 chars fits
    # exactly one of these two units.
    bounded = build_corpus_evidence(corpus_user, "", max_chars=150)
    assert 0 < len(bounded) <= 150
    assert "\n\n" not in bounded, "only the single unit that fits may be emitted"


def test_corpus_evidence_is_empty_for_a_user_with_no_corpus(corpus_user):
    """Honest degradation: no corpus means résumé-only evidence, never filler."""
    from app.services.evidence_corpus import build_corpus_evidence

    assert build_corpus_evidence(corpus_user, "Any job description") == ""


def test_load_corpus_file_rejects_a_shape_it_did_not_verify(tmp_path):
    from app.services.evidence_corpus import load_corpus_file

    path = tmp_path / "bad.json"
    path.write_text(json.dumps({"claims": "not a list"}), encoding="utf-8")
    with pytest.raises(ValueError):
        load_corpus_file(path)
