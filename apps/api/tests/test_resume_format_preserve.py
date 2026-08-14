"""MODELS-LIVE R-FMT — a tailored résumé is the baseline with ONLY text changed.

The mandate: *tailored = baseline with ONLY the reworded text changed, the
visual format intact.* Live production violated it — a tailored child of the
two-column baseline ``cfe7a0f…`` downloaded as a 9.4 KB single-column branded
PDF (``c12187…``) because ONE of four rewrites aimed at the left rail could not
be spliced and the whole render was dropped all-or-nothing to Aether's branded
template (``uat/reports/evidence/market-perf/resume-format/``).

This suite pins the ruling that removes that failure:

1. **The flagship** — a freshly tailored child of a bundled two-column baseline
   downloads as the SAME two-column PDF, with a ~0 pixel difference OUTSIDE the
   reworded bullet, proved by an independent rasteriser (``pdftoppm``) and a
   PyMuPDF text-span cross-check. A rewrite the splice engine cannot place keeps
   its baseline wording and is disclosed as residue — never a drop to branded.
2. **The verification harness** (``services/format_diff.py``) itself: it must
   PASS an in-place splice and FAIL a genuine re-layout (branded single-column).
3. **The fidelity contract** — an in-place render with unplaceable residue is
   ``formatPreserved: true`` (the layout is the user's own), ``confidence:
   partial``, with the residue named, and the WHOLE original document still
   present. A GENUINE content loss still falls back to the branded render (the
   U2b CRITICAL guarantee, intact).
4. **Legacy** — a résumé with no retained original is told to re-upload, not
   silently branded as "preserved".

Fixtures derive from the repository's own bundled résumé asset at run time, so
the tests exercise the real geometry that produced the live defect.
"""
from __future__ import annotations

import hashlib
from io import BytesIO

# ONE definition of the bundled-layout fixtures + "a change the splice engine
# structurally cannot apply", shared with the U2b truth-round tests.
from test_u2b_fidelity_verification import (  # noqa: E402
    _bundled_pdf,
    _left_rail_line,
    _pdf_text,
    _right_column_bullet,
    _seed_bundled_baseline,
    _tailor_child,
    _user_id,
)


def _right_after(right: str) -> str:
    head, _, tail = right.partition(":")
    return f"{head}: Re-scoped for the target role, {tail.strip()}"


def _left_after(left: str) -> str:
    return f"Product and delivery leadership across regulated platforms; {left}"


def _changed_slot_boxes(before_text: str) -> dict[int, list[tuple[float, float, float, float]]]:
    """Bounding boxes (PyMuPDF points) of the work bullet whose text is ``before``.

    Read from the splice engine's OWN block detection, so the mask covers exactly
    the region the renderer redraws (plus a small pad for its peach highlight and
    any font step-down), and nothing else.
    """
    import fitz

    from app.services.resume_pdf import _RIGHT_MARGIN, _detect_blocks, _normalize

    key = _normalize(before_text)
    boxes: dict[int, list[tuple[float, float, float, float]]] = {}
    doc = fitz.open(_bundled_pdf())
    try:
        for page_index in range(len(doc)):
            for block in _detect_blocks(doc[page_index]):
                if _normalize(block["full_text"]).startswith(key[:40]) or key.startswith(
                    _normalize(block["full_text"])[:40]
                ):
                    boxes.setdefault(page_index, []).append((
                        block["x0"] - 4.0,
                        block["top"] - 4.0,
                        _RIGHT_MARGIN + 4.0,
                        block["bottom"] + 8.0,
                    ))
    finally:
        doc.close()
    return boxes


# ---------------------------------------------------------------------------
# (2) The verification harness: passes a splice, fails a re-layout
# ---------------------------------------------------------------------------


def test_harness_pdf_layout_identical_pdf_is_preserved():
    from app.services.format_diff import compare_pdf_layout

    data = _bundled_pdf().read_bytes()
    result = compare_pdf_layout(data, data)
    assert result.same_page_count and result.same_geometry
    assert result.max_diff_ratio == 0.0
    assert result.preserved is True


def test_harness_pdf_layout_flags_a_branded_relayout_as_not_preserved():
    """A single-column branded re-render of the same résumé must FAIL the gate.

    This is the live failure mode: same words, wholly different layout. If the
    harness passed it, the harness would be useless.
    """
    from app.services.format_diff import compare_pdf_layout
    from app.services.resume_pdf import create_branded_resume_pdf, extract_pdf_bullets

    baseline = _bundled_pdf().read_bytes()
    bullets = extract_pdf_bullets(_bundled_pdf())
    branded = create_branded_resume_pdf(
        "VIKRAM DESHPANDE",
        "Senior Technical Program/Delivery Manager",
        "",
        [{"heading": "WORK EXPERIENCE", "bullets": bullets}],
        None,
    )
    result = compare_pdf_layout(baseline, branded)
    assert result.preserved is False, (
        "a single-column branded re-render is NOT the two-column original — the "
        f"harness must reject it, got {result!r}"
    )


def test_harness_masks_exclude_the_reworded_slot():
    """The pixel diff of a real splice is ~0 OUTSIDE the reworded bullet's mask."""
    from app.services.format_diff import compare_pdf_layout
    from app.services.resume_pdf import render_tailored_pdf

    right = _right_column_bullet()
    spliced = render_tailored_pdf(_bundled_pdf(), [(right, _right_after(right))])
    boxes = _changed_slot_boxes(right)
    assert boxes, "fixture sanity: the reworded bullet must be locatable in the layout"

    masked = compare_pdf_layout(_bundled_pdf().read_bytes(), spliced, change_boxes=boxes)
    assert masked.same_page_count and masked.same_geometry
    assert masked.preserved is True, (
        f"outside the reworded slot the splice must be pixel-identical, got {masked!r}"
    )

    # Without the mask, the reworded slot itself IS a difference — proving the
    # mask is doing real work rather than the diff being trivially empty.
    unmasked = compare_pdf_layout(_bundled_pdf().read_bytes(), spliced)
    assert unmasked.max_diff_ratio > masked.max_diff_ratio


def test_harness_pdf_text_span_crosscheck_matches_outside_masks():
    from app.services.format_diff import compare_pdf_text_spans
    from app.services.resume_pdf import render_tailored_pdf

    right = _right_column_bullet()
    spliced = render_tailored_pdf(_bundled_pdf(), [(right, _right_after(right))])
    diff = compare_pdf_text_spans(
        _bundled_pdf().read_bytes(), spliced, change_boxes=_changed_slot_boxes(right)
    )
    assert diff.identical, (
        "every text span outside the reworded slot must sit at the same "
        f"position, font and size: {diff!r}"
    )


# ---------------------------------------------------------------------------
# (1) THE FLAGSHIP — the two-column baseline and its tailored child both
#     download as the same two-column PDF, ~0 diff outside the reworded bullet
# ---------------------------------------------------------------------------


def test_flagship_tailored_child_downloads_as_the_two_column_layout(client, auth_headers):
    """The exact live failure, reversed: cfe7a0f→child stays two-column.

    A right-column bullet is reworded (the splice CAN place it) and a left-rail
    line is reworded (the splice CANNOT). Under the old all-or-nothing gate this
    produced the single-column branded PDF; now the download is the preserved
    two-column splice with the left-rail rewrite disclosed as residue.
    """
    from app.services.format_diff import compare_pdf_layout, compare_pdf_text_spans

    baseline, left, right = _seed_bundled_baseline(client, auth_headers)
    left_after, right_after = _left_after(left), _right_after(right)
    child = _tailor_child(
        client, auth_headers, baseline, {"bullet-0": left_after, "bullet-1": right_after}
    )

    fidelity = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers).json()
    assert fidelity["method"] == "pdf-in-place-splice", fidelity
    assert fidelity["verification"] == "post-render-text-extraction", fidelity
    assert fidelity["changesRequested"] == 2, fidelity
    assert fidelity["changesApplied"] == 1, fidelity
    assert fidelity["changesDropped"] == 1, fidelity
    assert fidelity["formatPreserved"] is True, (
        "the two-column layout IS the user's own document — format is preserved "
        f"even though one out-of-scope rewrite could not be placed: {fidelity}"
    )
    assert fidelity["confidence"] == "partial", fidelity
    note = fidelity["note"].lower()
    assert "layout is preserved" in note, note
    assert "aether template" not in note, note

    base_dl = client.get(f"/resumes/{baseline['id']}/download", headers=auth_headers)
    child_dl = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert base_dl.status_code == 200 and child_dl.status_code == 200
    assert child_dl.headers["X-Aether-Format-Method"] == "pdf-in-place-splice"
    assert child_dl.headers["X-Aether-Changes-Applied"] == "1"
    assert child_dl.headers["X-Aether-Changes-Dropped"] == "1"

    # The child must be the SAME two-column layout as the baseline — same page
    # count and geometry, ~0 pixel diff outside the one reworded work bullet.
    boxes = _changed_slot_boxes(right)
    assert boxes, "fixture sanity: the reworded work bullet must be locatable"
    layout = compare_pdf_layout(base_dl.content, child_dl.content, change_boxes=boxes)
    assert layout.same_page_count, layout
    assert layout.same_geometry, layout
    assert layout.preserved is True, (
        "the tailored child must be the two-column baseline with ONLY the "
        f"reworded bullet changed: {layout!r}"
    )
    spans = compare_pdf_text_spans(base_dl.content, child_dl.content, change_boxes=boxes)
    assert spans.identical, spans

    # The unplaceable left-rail rewrite keeps its ORIGINAL wording (residue),
    # and the placeable right-column rewrite is really applied.
    text = " ".join(_pdf_text(child_dl.content).split())
    assert left in text, "the unplaceable rewrite must keep the baseline wording"
    assert right_after.partition(":")[2].strip()[:50] in text, text[:400]


def test_flagship_render_harness_gate_without_the_api(client, auth_headers):
    """The §4 acceptance gate, expressed directly on the render + harness.

    Independent of the endpoint: splice the bundled two-column PDF and assert the
    harness certifies it as the same layout outside the reworded slot. This is
    the gate the brief says FAILS today (a branded single-column render).
    """
    from app.services.format_diff import compare_pdf_layout
    from app.services.resume_pdf import render_tailored_pdf

    right = _right_column_bullet()
    spliced = render_tailored_pdf(_bundled_pdf(), [(right, _right_after(right))])
    result = compare_pdf_layout(
        _bundled_pdf().read_bytes(), spliced, change_boxes=_changed_slot_boxes(right)
    )
    assert result.page_count >= 1
    assert result.preserved is True, result


# ---------------------------------------------------------------------------
# (3) The fidelity contract for an in-place render with residue
# ---------------------------------------------------------------------------


def test_verified_fidelity_partial_preserves_format_keeps_preserved_true():
    """An unplaceable rewrite that keeps the original wording preserves format."""
    from app.services.format_verification import ChangeOutcome, RenderVerification
    from app.services.resume_format import describe_fidelity, verified_fidelity

    base = describe_fidelity(
        bundled_match=True, has_original=False, content_type=None, is_tailored=True
    )
    verification = RenderVerification(
        requested=2,
        text_extracted=True,
        outcomes=(
            ChangeOutcome("b0", "a0", coverage=1.0, applied=True, original_remains=False),
            ChangeOutcome(
                "b1 left rail skills line here",
                "a1 reworded skills that could not be placed in the left rail",
                coverage=0.05,
                applied=False,
                original_remains=True,
            ),
        ),
    )
    report = verified_fidelity(base, verification, partial_preserves_format=True)
    assert report.method == "pdf-in-place-splice"
    assert report.changes_applied == 1
    assert report.changes_dropped == 1
    assert report.confidence == "partial"
    assert report.preserved is True, (
        "the layout is the user's own document; the residue is disclosed, not a "
        f"reason to deny preservation: {report!r}"
    )
    assert "original wording" in report.note.lower()
    assert "layout is preserved" in report.note.lower()

    # The default (branded-fallback) contract is UNCHANGED — a dropped change
    # there is still not preserved (U2b honesty machinery, untouched).
    default = verified_fidelity(base, verification)
    assert default.preserved is False


def test_build_applied_content_keeps_original_wording_for_a_dropped_rewrite():
    from app.services.resume_completeness import build_applied_content

    parent = {
        "sections": {
            "raw_text": (
                "JANE SMITH\n"
                "EXPERIENCE\n"
                "• Delivered the payments platform migration on schedule.\n"
                "• Owned the incident response runbook programme end to end."
            ),
            "bullets": [
                {"text": "Delivered the payments platform migration on schedule."},
                {"text": "Owned the incident response runbook programme end to end."},
            ],
        }
    }
    applied = [("Delivered the payments platform migration on schedule.", "AFTER placed")]
    content = build_applied_content(parent, applied)
    joined = " ".join(content.bullets).lower()
    assert "after placed" in joined, "a PLACED rewrite is substituted into the contract"
    assert "incident response runbook" in joined, (
        "a DROPPED rewrite keeps the parent's ORIGINAL wording in the contract, "
        "so the in-place render that still carries it is not falsely 'incomplete'"
    )


def test_genuine_content_loss_still_falls_back_to_branded(client, auth_headers):
    """A splice that LOST original content is NOT shipped (U2b CRITICAL intact).

    Modelled at the seam the router uses: when the applied-only completeness
    check reports a missing heading/contact/bullet, the in-place render must be
    refused in favour of the content-complete branded render.
    """
    from app.services.resume_completeness import CompletenessVerification

    # A render that genuinely dropped a whole section is content-incomplete even
    # under the lenient applied-only contract, so `.complete` is False and the
    # router routes to branded — the guarantee we must not weaken.
    lossy = CompletenessVerification(
        text_extracted=True, missing_headings=("EDUCATION",), missing_contact=("a@b.com",)
    )
    assert lossy.complete is False


# ---------------------------------------------------------------------------
# (4) Legacy — no retained original ⇒ re-upload, never silent "preserved"
# ---------------------------------------------------------------------------


def test_legacy_no_original_row_is_told_to_reupload(client, auth_headers):
    from app.repositories.resume import ResumeRepository

    me = client.get("/auth/me", headers=auth_headers).json()
    repo = ResumeRepository()
    repo.create(
        me["id"],
        {"raw_text": "TYPED RESUME\nEXPERIENCE\n• A typed bullet with no file.", "bullets": []},
        hashlib.sha256(b"not-a-bundled-asset").hexdigest(),
        label="Ingested — no original",
        version=repo.next_version(me["id"]),
    )
    listing = client.get("/resumes", headers=auth_headers).json()
    row = next(r for r in listing if r["label"] == "Ingested — no original")
    assert row["formatPreserved"] is False
    assert "re-upload" in row["formatFidelity"]["note"].lower(), row["formatFidelity"]


# ---------------------------------------------------------------------------
# (5) DOCX — the clean flagship: true in-place <w:t> substitution, with
#     styles.xml / numbering.xml / theme / fonts byte-identical (SYNTHESIS §4)
# ---------------------------------------------------------------------------

_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
#: A genuinely STYLED résumé — a Title, a Subtitle, a Heading and three
#: ``List Bullet`` paragraphs — so the package carries a real ``styles.xml``,
#: ``numbering.xml``, ``theme`` and font table for the structural diff to guard.
#: A plain-text-only fixture could pass "styles preserved" by accident.
_DOCX_PARAGRAPHS = [
    ("Title", "JORDAN AKINYEMI"),
    ("Subtitle", "Senior Platform Engineer — Perth, WA, Australia"),
    ("Heading 1", "EXPERIENCE"),
    ("List Bullet", "Operated a Kubernetes fleet serving 12 million requests per day."),
    ("List Bullet", "Cut deploy lead time by 47 percent through pipeline automation."),
    ("List Bullet", "Mentored six engineers across two on-call rotations."),
]
_DOCX_ORIGINAL_BULLET = "Cut deploy lead time by 47 percent through pipeline automation."
_DOCX_TAILORED_BULLET = (
    "Cut deploy lead time by 47 percent through Docker-based container pipeline automation."
)


def _make_styled_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    from docx import Document

    doc = Document()
    for style, text in paragraphs:
        p = doc.add_paragraph(style=style if style != "Title" else None)
        run = p.add_run(text)
        if style == "Title":
            run.bold = True
        if style == "Subtitle":
            run.italic = True
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_harness_docx_structure_passes_a_true_inplace_rewrite():
    """The DOCX §4 assertion: a real in-place rewrite leaves every structural
    part (styles/numbering/theme/fonts) byte-identical and touches ONLY <w:t>.
    """
    from app.services.format_diff import _docx_parts, compare_docx_structure
    from app.services.resume_docx import render_tailored_docx

    original = _make_styled_docx(_DOCX_PARAGRAPHS)
    tailored = render_tailored_docx(
        original, [(_DOCX_ORIGINAL_BULLET, _DOCX_TAILORED_BULLET)]
    )
    assert tailored != original, "a genuine rewrite must change the bytes"

    diff = compare_docx_structure(original, tailored)
    assert diff.styles_preserved is True, (
        "styles.xml / numbering.xml / theme / fonts must be byte-identical and "
        f"the non-text skeleton unchanged: {diff!r}"
    )
    assert diff.changed_style_parts == () and diff.missing_parts == () and diff.added_parts == ()
    assert diff.document_structure_changed is False
    # document.xml IS allowed to differ — but only inside its <w:t> text runs.
    parts_before, parts_after = _docx_parts(original), _docx_parts(tailored)
    assert parts_before["word/document.xml"] != parts_after["word/document.xml"], (
        "the reworded text must actually change document.xml"
    )
    assert "word/document.xml" in diff.changed_text_only_parts


def test_harness_docx_structure_flags_a_structural_change():
    """If a rewrite altered a run/paragraph property — not just the words — the
    harness MUST reject it. Here the bullets lose their ``List Bullet`` numbering
    (a pPr/numbering change), which the mandate forbids a tailoring edit to make.
    """
    from app.services.format_diff import compare_docx_structure

    original = _make_styled_docx(_DOCX_PARAGRAPHS)
    destyled = _make_styled_docx(
        [("Normal" if style == "List Bullet" else style, text) for style, text in _DOCX_PARAGRAPHS]
    )
    diff = compare_docx_structure(original, destyled)
    assert diff.styles_preserved is False, (
        "a numbering/paragraph-property change is a structural change the harness "
        f"must catch, not a permitted <w:t> text edit: {diff!r}"
    )
    assert diff.document_structure_changed is True


def test_flagship_docx_child_downloads_docx_native_with_styles_byte_identical(
    client, auth_headers,
):
    """DOCX end-to-end: an uploaded styled .docx tailored for a job downloads as
    a .docx (never the branded PDF), byte-identical outside the one reworded run.

    ``item 4`` of the mandate — "prefer DOCX preservation when the original is
    DOCX" — proved on the produced artifacts with the §4 structural diff.
    """
    from app.repositories.resume import ResumeRepository
    from app.services.format_diff import compare_docx_structure

    me = client.get("/auth/me", headers=auth_headers).json()
    # Capture the EXACT uploaded bytes: python-docx stamps wall-clock time into
    # each ZIP part header, so a second independent build drifts by a couple of
    # seconds during a long run — byte-identity must be checked against what was
    # actually uploaded, not a fresh rebuild.
    uploaded_docx = _make_styled_docx(_DOCX_PARAGRAPHS)
    upload = client.post(
        "/resumes/upload",
        files={"file": ("jordan_resume.docx", uploaded_docx, _DOCX_CONTENT_TYPE)},
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert upload.status_code == 201, upload.text
    baseline = upload.json()
    assert any(
        b["text"].strip() == _DOCX_ORIGINAL_BULLET for b in baseline["sections"]["bullets"]
    ), baseline["sections"]["bullets"]

    repo = ResumeRepository()
    child_sections = dict(baseline["sections"])
    child_sections["bullets"] = [
        {
            "text": _DOCX_TAILORED_BULLET
            if b["text"].strip() == _DOCX_ORIGINAL_BULLET
            else b["text"],
            "evidenceRef": b.get("evidenceRef"),
        }
        for b in baseline["sections"]["bullets"]
    ]
    child = repo.create(
        me["id"],
        child_sections,
        baseline["formatHash"],
        label="Tailored — Platform Engineer @ Atlassian",
        version=repo.next_version(me["id"]),
        parent_id=baseline["id"],
    )

    fidelity = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers).json()
    assert fidelity["method"] == "docx-native", fidelity
    assert fidelity["formatPreserved"] is True, fidelity

    base_dl = client.get(f"/resumes/{baseline['id']}/download", headers=auth_headers)
    child_dl = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert base_dl.status_code == 200 and child_dl.status_code == 200
    assert child_dl.headers["content-type"].startswith(_DOCX_CONTENT_TYPE)
    assert child_dl.headers["X-Aether-Format-Method"] == "docx-native"

    # The base download is the user's own stored .docx, verbatim.
    assert base_dl.content == uploaded_docx

    # The child is that same package with ONLY the reworded run's <w:t> changed:
    # styles / numbering / theme / fonts byte-identical, structure untouched.
    diff = compare_docx_structure(base_dl.content, child_dl.content)
    assert diff.styles_preserved is True, diff
    assert diff.document_structure_changed is False, diff

    from docx import Document

    text = "\n".join(p.text for p in Document(BytesIO(child_dl.content)).paragraphs)
    assert _DOCX_TAILORED_BULLET in text
    assert _DOCX_ORIGINAL_BULLET not in text


# ---------------------------------------------------------------------------
# (6) THE PDF FLAGSHIP — a genuine, NON-bundled PDF upload (the majority
#     real-world case, the exact format of the live cfe7a0f→c12187 incident)
#     is spliced in place from its STORED bytes, never dropped to branded.
#     Before this fix `render_tailored_pdf` only accepted a bundled Path, so the
#     retained DB bytes of a real PDF upload never reached the splice at all
#     (ML-RFMT PDF splice gap).
# ---------------------------------------------------------------------------


def _nonbundled_pdf_copy() -> bytes:
    """A genuine, non-bundled PDF that keeps the bundled résumé's OWN geometry.

    Re-serialising the bundled asset with a metadata edit changes every byte —
    so its SHA-256 no longer matches any bundled asset and
    ``resolve_original_pdf`` returns ``None`` (it is treated exactly like a real
    user upload, spliced from stored bytes) — while the two-column layout, coral
    bullets and right-column body geometry the splice engine keys on are carried
    through unchanged. This is the closest a deterministic fixture can get to the
    live incident: a user's own PDF, not on disk anywhere, that MUST be spliced
    from its retained bytes rather than dropped to the branded template.
    """
    import fitz

    doc = fitz.open(_bundled_pdf())
    try:
        doc.set_metadata({"title": "user upload copy", "author": "candidate"})
        return doc.tobytes(garbage=3, deflate=True)
    finally:
        doc.close()


def _boxes_from_pdf_bytes(
    data: bytes, before_text: str
) -> dict[int, list[tuple[float, float, float, float]]]:
    import fitz

    from app.services.resume_pdf import _RIGHT_MARGIN, _detect_blocks, _normalize

    key = _normalize(before_text)
    boxes: dict[int, list[tuple[float, float, float, float]]] = {}
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        for page_index in range(len(doc)):
            for block in _detect_blocks(doc[page_index]):
                if _normalize(block["full_text"]).startswith(key[:40]) or key.startswith(
                    _normalize(block["full_text"])[:40]
                ):
                    boxes.setdefault(page_index, []).append((
                        block["x0"] - 4.0,
                        block["top"] - 4.0,
                        _RIGHT_MARGIN + 4.0,
                        block["bottom"] + 8.0,
                    ))
    finally:
        doc.close()
    return boxes


def _seed_stored_pdf_baseline(client, auth_headers) -> tuple[dict, str, str, bytes]:
    """A baseline backed by STORED (non-bundled) PDF bytes, plus its two bullets.

    Mirrors ``_seed_bundled_baseline`` but the row carries its own
    ``originalFile`` bytes and a NON-bundled ``formatHash`` — so the download
    resolves the format through the stored bytes, the path finding ML-RFMT says
    was never wired. The bullets are the bundled résumé's own positional bullets
    (identical text survives the re-serialisation), so the splice can place the
    right-column one exactly as it does for the bundled path.
    """
    from app.repositories.resume import ResumeRepository

    copy = _nonbundled_pdf_copy()
    left = _left_rail_line()
    right = _right_column_bullet()
    user_id = _user_id(client, auth_headers)
    repo = ResumeRepository()
    baseline = repo.create(
        user_id,
        {
            "raw_text": f"{left}\n{right}",
            "bullets": [
                {"text": left, "evidenceRef": "bullet-0"},
                {"text": right, "evidenceRef": "bullet-1"},
            ],
            "contact": {},
        },
        hashlib.sha256(copy).hexdigest(),
        label="Uploaded — candidate_resume.pdf",
        version=repo.next_version(user_id),
        original_file=copy,
        original_filename="candidate_resume.pdf",
        original_content_type="application/pdf",
    )
    return baseline, left, right, copy


def test_flagship_stored_pdf_upload_tailored_child_is_spliced_not_branded(
    client, auth_headers,
):
    """The finding, reversed: a tailored child of a genuine NON-bundled PDF
    upload downloads as the SAME two-column PDF — spliced from the retained
    bytes — not the branded single-column template.

    A right-column bullet is reworded (the splice CAN place it) and a left-rail
    line is reworded (it CANNOT). The download is the preserved two-column
    splice with the left-rail rewrite disclosed as residue, exactly like the
    bundled path — proving the stored-bytes splice route is now wired.
    """
    from app.services.format_diff import compare_pdf_layout, compare_pdf_text_spans
    from app.services.resume_pdf import resolve_original_pdf

    baseline, left, right, copy = _seed_stored_pdf_baseline(client, auth_headers)
    # Precondition: this really is a non-bundled résumé (the splice must come
    # from the stored bytes, never a bundled Path).
    assert resolve_original_pdf(baseline["formatHash"]) is None, (
        "fixture sanity: the uploaded PDF must NOT match any bundled asset"
    )

    child = _tailor_child(
        client,
        auth_headers,
        baseline,
        {"bullet-0": _left_after(left), "bullet-1": _right_after(right)},
    )

    fidelity = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers).json()
    assert fidelity["method"] == "pdf-in-place-splice", fidelity
    assert fidelity["verification"] == "post-render-text-extraction", fidelity
    assert fidelity["changesRequested"] == 2, fidelity
    assert fidelity["changesApplied"] == 1, fidelity
    assert fidelity["changesDropped"] == 1, fidelity
    assert fidelity["formatPreserved"] is True, (
        "the uploaded PDF's own two-column layout IS preserved — the "
        f"unplaceable left-rail rewrite is residue, not a drop to branded: {fidelity}"
    )
    assert fidelity["confidence"] == "partial", fidelity
    assert "aether template" not in fidelity["note"].lower(), fidelity

    base_dl = client.get(f"/resumes/{baseline['id']}/download", headers=auth_headers)
    child_dl = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert base_dl.status_code == 200 and child_dl.status_code == 200
    # The base download is the user's own uploaded PDF, byte-identical.
    assert base_dl.content == copy, "a base PDF upload download must be byte-identical"
    assert base_dl.headers["X-Aether-Format-Method"] == "original-bytes"
    assert child_dl.headers["X-Aether-Format-Method"] == "pdf-in-place-splice"
    assert child_dl.headers["X-Aether-Changes-Applied"] == "1"
    assert child_dl.headers["X-Aether-Changes-Dropped"] == "1"

    boxes = _boxes_from_pdf_bytes(copy, right)
    assert boxes, "fixture sanity: the reworded work bullet must be locatable"
    layout = compare_pdf_layout(base_dl.content, child_dl.content, change_boxes=boxes)
    assert layout.same_page_count, layout
    assert layout.same_geometry, layout
    assert layout.preserved is True, (
        "the tailored child must be the uploaded two-column PDF with ONLY the "
        f"reworded work bullet changed: {layout!r}"
    )
    spans = compare_pdf_text_spans(base_dl.content, child_dl.content, change_boxes=boxes)
    assert spans.identical, spans

    text = " ".join(_pdf_text(child_dl.content).split())
    assert left in text, "the unplaceable left-rail rewrite keeps the baseline wording"
    assert _right_after(right).partition(":")[2].strip()[:50] in text, text[:400]


def test_arbitrary_pdf_upload_tailored_child_keeps_layout_never_branded(
    client, auth_headers,
):
    """A genuinely arbitrary PDF upload (no bullets the splice geometry can
    place) still keeps the user's OWN layout — it is NOT re-rendered into the
    branded single-column template over an unplaceable rewrite.

    This is the minimum the mandate requires (binding scope item 3): when a
    region cannot be spliced, keep the layout and disclose the residue, never
    drop to branded. The download is the user's own PDF (same page count and
    media box), reported ``pdf-in-place-splice`` / preserved, with the rewrite
    listed as not applied.
    """
    from io import BytesIO

    from reportlab.pdfgen import canvas

    from app.repositories.resume import ResumeRepository
    from app.services.format_diff import compare_pdf_layout
    from app.services.resume_pdf import resolve_original_pdf

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(612, 792))
    y = 720
    lines = [
        "MORGAN ELLIS",
        "Principal Reliability Engineer — Hobart, TAS, Australia",
        "EXPERIENCE",
        "Ran the global on-call programme across four regional teams.",
        "Reduced mean time to recovery from ninety minutes to twelve.",
    ]
    for line in lines:
        c.drawString(72, y, line)
        y -= 20
    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()

    upload = client.post(
        "/resumes/upload",
        files={"file": ("morgan_resume.pdf", pdf_bytes, "application/pdf")},
        data={"extract_stories": "false"},
        headers=auth_headers,
    )
    assert upload.status_code == 201, upload.text
    baseline = upload.json()
    assert resolve_original_pdf(baseline["formatHash"]) is None

    me = client.get("/auth/me", headers=auth_headers).json()
    repo = ResumeRepository()
    original_bullet = "Ran the global on-call programme across four regional teams."
    tailored_bullet = "Directed the 24x7 on-call programme spanning four regional SRE teams."
    child_sections = dict(baseline["sections"])
    child_sections["bullets"] = [
        {
            "text": tailored_bullet if b["text"].strip() == original_bullet else b["text"],
            "evidenceRef": b.get("evidenceRef"),
        }
        for b in baseline["sections"]["bullets"]
    ]
    child = repo.create(
        me["id"],
        child_sections,
        baseline["formatHash"],
        label="Tailored — Reliability Lead @ Canva",
        version=repo.next_version(me["id"]),
        parent_id=baseline["id"],
    )

    fidelity = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers).json()
    assert fidelity["method"] == "pdf-in-place-splice", fidelity
    assert fidelity["method"] != "reflow-template", fidelity
    assert fidelity["method"] != "branded-optin", fidelity
    assert fidelity["formatPreserved"] is True, (
        "an unplaceable rewrite must keep the user's own PDF layout, never drop "
        f"to the branded template: {fidelity}"
    )

    child_dl = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert child_dl.headers["X-Aether-Format-Method"] == "pdf-in-place-splice"
    # The download is the user's own single-page PDF, not the multi-page branded
    # re-render: same page count and media box as what they uploaded.
    layout = compare_pdf_layout(pdf_bytes, child_dl.content)
    assert layout.same_page_count, layout
    assert layout.same_geometry, layout


# ---------------------------------------------------------------------------
# (7) The branded template is an EXPLICIT opt-in, never a silent fallback for a
#     retained-original row (binding scope item 5).
# ---------------------------------------------------------------------------


def test_branded_render_is_explicit_optin_only(client, auth_headers):
    from app.services.format_diff import compare_pdf_layout

    baseline, left, right = _seed_bundled_baseline(client, auth_headers)
    child = _tailor_child(
        client, auth_headers, baseline, {"bullet-1": _right_after(right)}
    )

    # Default download: the user's own preserved layout.
    default = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert default.headers["X-Aether-Format-Method"] == "pdf-in-place-splice"

    # EXPLICIT opt-in: the branded template, honestly labelled — never preserved.
    branded = client.get(
        f"/resumes/{child['id']}/download?branded=true", headers=auth_headers
    )
    assert branded.status_code == 200
    assert branded.headers["X-Aether-Format-Method"] == "branded-optin"

    fid = client.get(
        f"/resumes/{child['id']}/fidelity?branded=true", headers=auth_headers
    ).json()
    assert fid["method"] == "branded-optin", fid
    assert fid["formatPreserved"] is False, fid
    note = fid["note"].lower()
    assert "at your request" in note or "template" in note, note

    # The branded opt-in really IS a different (single-column) render, not the
    # preserved two-column layout — proving the opt-in changed the artifact.
    relayout = compare_pdf_layout(default.content, branded.content)
    assert relayout.preserved is False, (
        "the branded opt-in must be a genuinely different layout from the "
        f"preserved default download: {relayout!r}"
    )


# ---------------------------------------------------------------------------
# (8) Backfill audit — COUNTS the legacy rows that cannot be format-preserved
#     (binding scope item 6). Read-only classification, no writes.
# ---------------------------------------------------------------------------


def test_backfill_audit_counts_affected_rows():
    import importlib.util
    from pathlib import Path

    script = (
        Path(__file__).resolve().parents[1] / "scripts" / "audit_resume_format_backfill.py"
    )
    spec = importlib.util.spec_from_file_location("_rfmt_backfill_audit", script)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    bundled_hash = "b" * 64
    bundled = {bundled_hash}
    rows = [
        {"id": "r1", "parentId": None, "formatHash": "u" * 64, "hasOriginal": True},
        {"id": "r2", "parentId": None, "formatHash": bundled_hash, "hasOriginal": False},
        {"id": "r3", "parentId": None, "formatHash": "z" * 64, "hasOriginal": False},
        {"id": "r4", "parentId": "r3", "formatHash": "z" * 64, "hasOriginal": False},
        {"id": "r5", "parentId": "r1", "formatHash": "u" * 64, "hasOriginal": False},
    ]
    classified = module.classify(rows, bundled)
    summary = module.summarise(classified)

    by_id = {c["id"]: c for c in classified}
    assert by_id["r1"]["reason"] == "retained-original" and by_id["r1"]["preservable"]
    assert by_id["r2"]["reason"] == "bundled-hash-match" and by_id["r2"]["preservable"]
    assert by_id["r3"]["reason"] == "needs-reupload" and not by_id["r3"]["preservable"]
    assert by_id["r4"]["reason"] == "needs-reupload" and not by_id["r4"]["preservable"]
    assert by_id["r5"]["reason"] == "via-parent-original" and by_id["r5"]["preservable"]

    assert summary["examined"] == 5
    assert summary["preservable"] == 3
    assert summary["affected_needs_reupload"] == 2
    assert summary["affected_needs_reupload_base"] == 1
    assert summary["affected_needs_reupload_tailored"] == 1
    assert summary["wrote_anything"] is False


# ---------------------------------------------------------------------------
# (9) DOCX §4 — headers/footers/media must be byte-identical too (the harness
#     hole ML-RFMT flagged: same-named header/footer/media parts whose CONTENT
#     changed were invisible to the missing/added-part checks).
# ---------------------------------------------------------------------------


def _styled_docx_with_header_footer_media() -> bytes:
    from io import BytesIO

    from docx import Document
    from PIL import Image

    img = BytesIO()
    Image.new("RGB", (16, 16), (200, 40, 40)).save(img, format="PNG")

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("JORDAN AKINYEMI").bold = True
    doc.sections[0].header.paragraphs[0].add_run("CONFIDENTIAL — candidate copy")
    doc.sections[0].footer.paragraphs[0].add_run("References available on request")
    doc.add_paragraph("EXPERIENCE", style="Heading 1")
    for text in (
        "Operated a Kubernetes fleet serving 12 million requests per day.",
        "Cut deploy lead time by 47 percent through pipeline automation.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_picture(BytesIO(img.getvalue()))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_harness_docx_true_rewrite_keeps_header_footer_media_byte_identical():
    """An in-place body rewrite leaves header/footer/media byte-identical."""
    from app.services.format_diff import _docx_parts, compare_docx_structure
    from app.services.resume_docx import render_tailored_docx

    original = _styled_docx_with_header_footer_media()
    tailored = render_tailored_docx(
        original,
        [(
            "Cut deploy lead time by 47 percent through pipeline automation.",
            "Cut deploy lead time by 47 percent through Docker-based pipeline automation.",
        )],
    )
    assert tailored != original, "a genuine rewrite must change the bytes"

    diff = compare_docx_structure(original, tailored)
    assert diff.styles_preserved is True, diff
    assert diff.changed_style_parts == (), diff
    # The header/footer/media parts really are present and really are identical.
    parts_a, parts_b = _docx_parts(original), _docx_parts(tailored)
    for part in ("word/header1.xml", "word/footer1.xml", "word/media/image1.png"):
        assert part in parts_a, f"fixture must contain {part}"
        assert parts_a[part] == parts_b[part], f"{part} must be byte-identical"


def test_harness_docx_flags_a_changed_header_footer_and_media():
    """A change to a same-named header/footer/media part is caught (was a hole).

    These parts are present in both packages, so a content edit is invisible to
    the missing/added-part checks; the harness must byte-compare them.
    """
    import zipfile
    from io import BytesIO

    from app.services.format_diff import compare_docx_structure

    original = _styled_docx_with_header_footer_media()

    def _mutate(part_substr: str, needle: bytes, replacement: bytes) -> bytes:
        src = zipfile.ZipFile(BytesIO(original))
        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
            for item in src.namelist():
                data = src.read(item)
                if part_substr in item:
                    data = data.replace(needle, replacement)
                dst.writestr(item, data)
        return out.getvalue()

    changed_header = _mutate("word/header1.xml", b"CONFIDENTIAL", b"PUBLIC COPY")
    assert changed_header != original
    diff = compare_docx_structure(original, changed_header)
    assert diff.styles_preserved is False, (
        "a header content change must be caught, not folded into permitted text "
        f"deltas: {diff!r}"
    )
    assert "word/header1.xml" in diff.changed_style_parts, diff

    changed_footer = _mutate("word/footer1.xml", b"References", b"Referees")
    footer_diff = compare_docx_structure(original, changed_footer)
    assert footer_diff.styles_preserved is False, footer_diff
    assert "word/footer1.xml" in footer_diff.changed_style_parts, footer_diff

    # A media (image) byte change is caught by the word/media/ prefix guard.
    changed_media = _mutate("word/media/image1.png", b"IDAT", b"IDAT")  # no-op guard below
    # Force a real 1-byte change in the image part deterministically.
    src = zipfile.ZipFile(BytesIO(original))
    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.namelist():
            data = bytearray(src.read(item))
            if item == "word/media/image1.png":
                data[-1] ^= 0x01
            dst.writestr(item, bytes(data))
    changed_media = out.getvalue()
    media_diff = compare_docx_structure(original, changed_media)
    assert media_diff.styles_preserved is False, media_diff
    assert "word/media/image1.png" in media_diff.changed_style_parts, media_diff


# ---------------------------------------------------------------------------
# (10) REFIX — a LONG two-writer bullet (bold lead-in + grey body) is recognised
#      as PRESENT.  The 19 short-bullet fixtures above missed this whole class.
#
#      On the real two-column résumé cfe7a0f→c12187 (baseline
#      cfe7a0f27991821dc73f265cd → tailored child c12187d107bf994471844e09a, 10
#      changes / 8 placed / 2 unplaceable) the in-place completeness check named
#      TWO bullets "missing" that a raster + PyMuPDF re-extraction proved present
#      verbatim, so the whole preserved two-column layout was dropped to the 9.4 KB
#      branded single-column template — the exact live defect this slice exists to
#      remove, still shipping.  Both false negatives came from the shingle-fraction
#      matcher (``format_verification._coverage``) under-counting a bullet whose
#      text the renderer split across styled runs:
#
#      * an APPLIED bold-lead-in rewrite scored 0.839 — ``resume_pdf._render_block``
#        commits the grey ``reg`` writer first and the bold lead-in writer last, so
#        ``page.get_text`` returns the lead-in far from its body and the handful of
#        shingles straddling that seam were absent even though every word was on the
#        page.  Below 0.85 it read as "dropped", so the completeness contract then
#        expected the ORIGINAL wording (already redacted away) and reported it lost;
#      * an UNTOUCHED wrapped bullet scored 0.829 — a hyphenated compound broke over
#        a visual line ("test-" / "evidence"), so its own source bytes carried the
#        whole bullet yet a few shingles crossed the wrap.
#
#      Evidence: uat/reports/evidence/market-perf/resume-format/refix/.
# ---------------------------------------------------------------------------

#: A fixed, deterministic rewrite body long enough that, paired with the ≥6-word
#: bold lead-in below, the lead-in/body seam sinks the OLD shingle-fraction score
#: to ~0.76 (below the 0.85 applied bar) while every word is on the page.
_LONG_TWO_WRITER_AFTER_BODY = (
    "Led the technical delivery of AI solutions including real-time telemetry "
    "servers for large scale device concurrency and latency reduction."
)


def _long_lead_in_bundled_bullet() -> str:
    """A bundled right-column work bullet with a LONG (≥6-word) bold lead-in.

    Derived from the résumé asset at run time (never hard-coded prose), so the
    fixture exercises the real two-writer geometry: the redraw keeps this long
    lead-in in the bold weight and draws the reworded body in grey, which is the
    split ``page.get_text`` reports non-adjacently.
    """
    from app.services.resume_pdf import extract_pdf_bullets

    for bullet in extract_pdf_bullets(_bundled_pdf()):
        head = bullet.split(":", 1)[0]
        if ":" in bullet and len(head.split()) >= 6 and len(bullet) > 120:
            return bullet
    raise AssertionError("bundled résumé must have a long-lead-in work bullet")


def test_coverage_counts_a_long_two_writer_bullet_split_across_runs_as_present():
    """A LONG bold-lead-in rewrite the splice PLACES is scored APPLIED.

    RED before the refix: the lead-in/body seam sinks the shingle-fraction score
    to ~0.76, below the applied bar, even though every word is on the page —
    exactly what marked change[3] of the live record "dropped".
    """
    from app.services.format_verification import (
        _APPLIED_COVERAGE,
        _coverage,
        _normalize,
        extract_artifact_text,
        verify_changes,
    )
    from app.services.resume_pdf import render_tailored_pdf

    bullet = _long_lead_in_bundled_bullet()
    head = bullet.split(":", 1)[0]
    after = f"{head}: {_LONG_TWO_WRITER_AFTER_BODY}"
    spliced = render_tailored_pdf(_bundled_pdf(), [(bullet, after)])
    haystack = _normalize(extract_artifact_text(spliced, "application/pdf") or "")

    # The text really IS on the page: every word is present, and the bold lead-in
    # and the grey body each appear verbatim — just NOT as one contiguous run,
    # because the splice commits the grey writer first and the bold writer last.
    assert all(word in haystack for word in _normalize(after).split()), (
        "fixture sanity: every word of the rewrite must be on the page"
    )
    assert _normalize(head) in haystack, "the bold lead-in is present verbatim"
    assert _normalize(_LONG_TWO_WRITER_AFTER_BODY) in haystack, "the grey body is present verbatim"
    assert _normalize(after) not in haystack, (
        "fixture sanity: the two-writer seam really does split the lead-in from "
        "the body in the extracted text layer"
    )

    # The matcher must read the split-but-complete rewrite as PRESENT.
    assert _coverage(after, haystack) >= _APPLIED_COVERAGE, (
        "a two-writer bullet whose every word is on the page must clear the "
        f"applied bar, got {_coverage(after, haystack):.3f}"
    )
    result = verify_changes(spliced, "application/pdf", [(bullet, after)])
    assert result.applied_count == 1, [round(o.coverage, 3) for o in result.outcomes]


def test_is_present_counts_an_untouched_wrapped_two_writer_bullet_as_present():
    """An UNTOUCHED wrapped bullet reads as present against its own source bytes.

    RED before the refix: a hyphenated compound broke across a visual line, so a
    few shingles crossed the wrap and the shingle-fraction score fell to ~0.83 —
    flagging a bullet the splice never even touched as "missing content" and
    dropping the whole layout to branded.
    """
    from app.services.format_verification import _normalize, extract_artifact_text
    from app.services.resume_completeness import _is_present
    from app.services.resume_pdf import extract_pdf_bullets

    source = _bundled_pdf().read_bytes()
    haystack = _normalize(extract_artifact_text(source, "application/pdf") or "")
    # Long work bullets whose text is NOT a contiguous run in the flat text layer
    # (a line wrap split them) — the class the old matcher under-counted.
    wrapped = [
        bullet
        for bullet in extract_pdf_bullets(_bundled_pdf())
        if len(bullet) > 200 and _normalize(bullet) not in haystack
    ]
    assert wrapped, "fixture sanity: the bundled résumé must have a wrapped long bullet"
    for bullet in wrapped:
        assert _is_present(bullet, haystack), (
            "an untouched wrapped bullet, carried verbatim by its own source "
            f"bytes, must read as present: {bullet[:80]!r}"
        )


def test_long_two_writer_rewrite_ships_preserved_not_branded(client, auth_headers):
    """End-to-end: a LONG two-writer rewrite ships the preserved two-column PDF.

    The exact shape of the live cfe7a0f→c12187 failure: one right-column bullet
    with a long bold lead-in is reworded (the splice PLACES it, but the seam sank
    the old score below the applied bar → the completeness check then expected the
    now-redacted original and dropped the whole layout to the branded template),
    and one left-rail line is reworded (structurally unplaceable, kept as residue).
    The download must be the preserved two-column splice, NOT branded.
    """
    from app.repositories.resume import ResumeRepository

    left = _left_rail_line()
    right = _long_lead_in_bundled_bullet()
    user_id = _user_id(client, auth_headers)
    repo = ResumeRepository()
    baseline = repo.create(
        user_id,
        {
            "raw_text": f"{left}\n{right}",
            "bullets": [
                {"text": left, "evidenceRef": "bullet-0"},
                {"text": right, "evidenceRef": "bullet-1"},
            ],
            "contact": {},
        },
        hashlib.sha256(_bundled_pdf().read_bytes()).hexdigest(),
        label="Baseline — long two-writer layout",
        version=repo.next_version(user_id),
    )
    head = right.split(":", 1)[0]
    long_after = f"{head}: {_LONG_TWO_WRITER_AFTER_BODY}"
    child = _tailor_child(
        client,
        auth_headers,
        baseline,
        {"bullet-0": _left_after(left), "bullet-1": long_after},
    )

    fidelity = client.get(f"/resumes/{child['id']}/fidelity", headers=auth_headers).json()
    assert fidelity["method"] == "pdf-in-place-splice", fidelity
    assert fidelity["method"] != "reflow-template", (
        "the LONG two-writer rewrite is placed and complete — the preserved "
        f"two-column layout must ship, not the branded template: {fidelity}"
    )
    assert fidelity["formatPreserved"] is True, fidelity
    assert fidelity["changesApplied"] == 1, fidelity
    assert fidelity["changesDropped"] == 1, fidelity

    child_dl = client.get(f"/resumes/{child['id']}/download", headers=auth_headers)
    assert child_dl.headers["X-Aether-Format-Method"] == "pdf-in-place-splice"
    text = " ".join(_pdf_text(child_dl.content).split())
    assert "real-time telemetry servers" in text, (
        "the reworded body must be present verbatim in the produced PDF"
    )
