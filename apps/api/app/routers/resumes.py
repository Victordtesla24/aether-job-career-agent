"""Resumes router — versioned resume access + diff (P2-S05)."""
from __future__ import annotations

import hashlib
import os
import urllib.parse
import zipfile
from io import BytesIO
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, Response, UploadFile, status
from pydantic import BaseModel, Field

from app.middleware.auth import CurrentUser
from app.repositories.resume import ResumeRepository

router = APIRouter()

#: Hard ceiling on an uploaded résumé (U2a). The bytes are now persisted whole
#: in ``Resume.originalFile``, so an unbounded upload would be an unbounded row;
#: 10MB is far above any real résumé (the bundled reference PDFs are ~100KB)
#: while still refusing an accidental or hostile large-file POST. Enforced on a
#: BOUNDED read, so an oversized body is never fully buffered.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

_PDF_MAGIC = b"%PDF"
#: Local-file-header / empty-archive / spanned-archive ZIP signatures. A .docx
#: is an OOXML ZIP, so this is the first (cheap) gate before opening it.
_ZIP_MAGICS = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")

_PDF_CONTENT_TYPE = "application/pdf"
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_TEXT_EXTENSIONS = (".txt", ".text", ".md", ".markdown")

#: Honest rejection copy (MON-012). Before U2a, ANY non-PDF upload was decoded
#: with ``errors="replace"``, so a .docx, a screenshot or a corrupt file was
#: silently persisted as a Resume row full of U+FFFD replacement characters and
#: presented to the user as their résumé. Naming the formats Aether genuinely
#: supports is the honest answer; guessing is not.
_UNSUPPORTED_FORMAT_DETAIL = (
    "Unsupported file format. Aether reads PDF (.pdf), Word (.docx) and "
    "plain-text (.txt/.md) résumés; this file is not a readable document in "
    "any of those formats."
)


@router.get("")
def list_resumes(current_user: CurrentUser) -> list[dict[str, Any]]:
    return _with_format_preserved(ResumeRepository().list_by_user(current_user["id"]))


def _with_format_preserved(resumes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Stamp each résumé with an honest ``formatPreserved`` boolean (MON-011).

    ``True`` ONLY when ``GET /resumes/{id}/download`` would genuinely reproduce
    the original document — i.e. when ``resolve_original_pdf`` finds a bundled
    asset whose digest matches, which is the exact condition that endpoint
    branches on, resolved through the SAME parent-then-self ``formatHash``
    precedence it uses. Every other résumé (every real user upload, whose
    ``formatHash`` is a digest of the user's OWN bytes/text and so can never
    collide with a bundled asset) downloads as the re-flowed branded template,
    and now says so.

    Before this, the Resume Studio "Format Integrity Check" panel inferred a
    preservation claim from ``formatHash === baseHash`` — a self-comparison
    that is trivially true for a base résumé and says nothing about the
    download path — so every paying user was told their typography, spacing,
    columns and margins were preserved for a file that re-flows.

    The parent lookup reads the SAME list (a résumé's parent is another of that
    user's own versions), so this adds no query.
    """
    from app.services.resume_pdf import bundled_format_hashes

    bundled = bundled_format_hashes()
    by_id = {resume["id"]: resume for resume in resumes}
    stamped: list[dict[str, Any]] = []
    for resume in resumes:
        parent = by_id.get(resume.get("parentId"))
        format_hash = (parent or resume).get("formatHash") or resume.get("formatHash")
        stamped.append({
            **resume,
            "formatPreserved": bool(format_hash) and format_hash in bundled,
        })
    return stamped


class ResumeIngestRequest(BaseModel):
    """Register an additional root resume for the authenticated user.

    Used to ingest alternate resume variants (e.g. the BA-positioned resume)
    so they live in the database, appear in Resume Studio, and are selectable
    for tailoring runs. ``raw_text`` is the full extracted resume text —
    bullets are derived server-side so the anti-fabrication evidence index
    stays consistent with the tailoring service.
    """

    label: str = Field(min_length=1, max_length=120)
    raw_text: str = Field(min_length=50)
    contact: dict[str, Any] | None = None
    format_hash: str | None = Field(default=None, max_length=64)


@router.post("", status_code=status.HTTP_201_CREATED)
def create_resume(body: ResumeIngestRequest, current_user: CurrentUser) -> dict[str, Any]:
    """Ingest a new root resume version (Phase-2 audit Section C)."""
    from app.services.resume_tailor import extract_bullets

    sections = {
        "raw_text": body.raw_text,
        "bullets": [
            {"text": b, "evidenceRef": f"bullet-{i}"}
            for i, b in enumerate(extract_bullets(body.raw_text))
        ],
        "contact": body.contact or {},
    }
    format_hash = body.format_hash or hashlib.sha256(body.raw_text.encode()).hexdigest()[:16]
    repo = ResumeRepository()
    return repo.create(
        current_user["id"],
        sections,
        format_hash,
        label=body.label,
        version=repo.next_version(current_user["id"]),
    )


def _looks_like_docx(data: bytes) -> bool:
    """True when ``data`` really is an OOXML word-processing package.

    Checked by CONTENT, never by the client's filename or Content-Type: the
    bytes must be a readable ZIP that contains the WordprocessingML part
    (``word/document.xml``). Random bytes behind a ``.docx`` extension fail
    here and are rejected honestly instead of being decoded into garbage text.
    """
    if not data.startswith(_ZIP_MAGICS):
        return False
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            return any(name.startswith("word/") for name in archive.namelist())
    except (zipfile.BadZipFile, OSError):
        return False


def _extract_pdf_text(data: bytes) -> str:
    import fitz

    try:
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception as exc:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY, f"Could not parse PDF: {exc}"
        ) from exc


def _extract_docx_text(data: bytes) -> str:
    """Real WordprocessingML text extraction (U2a / R-F3).

    Reads the document's own paragraph and table model via ``python-docx``, so
    a .docx yields the actual words the user wrote. Paragraph order and blank
    lines are preserved because ``extract_bullets`` uses line structure (marker
    lines, all-caps section banners) to reassemble bullets. Table cells are
    appended after the body paragraphs — ``document.paragraphs`` covers only
    body-level paragraphs, so nothing is emitted twice.

    A damaged package fails HONESTLY with a 422, exactly like the sibling
    ``_extract_pdf_text``. A corrupt .docx surfaces low-level errors from every
    layer python-docx sits on, and they share no useful base class: a malformed
    ``[Content_Types].xml`` raises ``lxml.etree.XMLSyntaxError`` (a
    ``SyntaxError``), a damaged deflate stream inside ``word/document.xml``
    raises ``zlib.error``. Enumerating that family is a losing game and every
    miss is an unhandled 500 instead of the honest rejection MON-012 promises,
    so any parse failure is reported as an unprocessable entity.
    """
    from docx import Document

    try:
        document = Document(BytesIO(data))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
    except Exception as exc:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY, f"Could not parse DOCX: {exc}"
        ) from exc
    return "\n".join(lines)


def _extract_upload_text(filename: str, data: bytes) -> tuple[str, str]:
    """Extract résumé text from an upload, and report the format Aether VERIFIED.

    Returns ``(raw_text, content_type)`` where ``content_type`` is derived from
    what the bytes actually are (or, for plain text, from the extension) — never
    echoed from the client's Content-Type header, which is unverified input.
    Anything that is not a readable PDF, DOCX or UTF-8 text document raises an
    honest 422 (MON-012) rather than being force-decoded into noise.
    """
    lowered = filename.lower()
    if data.startswith(_PDF_MAGIC) or lowered.endswith(".pdf"):
        return _extract_pdf_text(data), _PDF_CONTENT_TYPE
    if _looks_like_docx(data):
        return _extract_docx_text(data), _DOCX_CONTENT_TYPE
    if lowered.endswith(_TEXT_EXTENSIONS):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(
                status.HTTP_422_UNPROCESSABLE_ENTITY,
                "This file is not valid UTF-8 text, so it cannot be read as a "
                "plain-text résumé. Supported formats: PDF (.pdf), Word (.docx) "
                "and UTF-8 text (.txt/.md).",
            ) from exc
        content_type = (
            "text/markdown; charset=utf-8"
            if lowered.endswith((".md", ".markdown"))
            else "text/plain; charset=utf-8"
        )
        return text, content_type
    raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, _UNSUPPORTED_FORMAT_DETAIL)


def _safe_upload_filename(filename: str) -> str:
    """The upload's own name, stripped of anything that is not a file name.

    Drops directory components (a browser or a hostile client may send a path),
    non-printable characters and quotes — the same characters that would let a
    stored name break out of the ``Content-Disposition`` header when it is
    served back by ``GET /resumes/{id}/original``.
    """
    base = os.path.basename(filename.replace("\\", "/")).strip()
    cleaned = "".join(ch for ch in base if ch.isprintable() and ch not in '"\\')
    return cleaned[:255] or "resume"


def _content_disposition(filename: str) -> str:
    """An ``attachment`` disposition carrying the original file name.

    Non-latin-1 names cannot travel in a bare ``filename=`` parameter, so they
    are emitted as RFC 5987 ``filename*=UTF-8''…`` instead of being mangled.
    """
    try:
        filename.encode("latin-1")
    except UnicodeEncodeError:
        quoted = urllib.parse.quote(filename, safe="")
        return f"attachment; filename*=UTF-8''{quoted}"
    return f'attachment; filename="{filename}"'


@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_resume(
    current_user: CurrentUser,
    file: UploadFile = File(...),
    extract_stories: bool = Form(default=False),
) -> dict[str, Any]:
    """Upload a resume file as a new root version (SC-ST-03).

    Extracts text server-side and registers a new ROOT resume through the same
    section-building path as JSON ingestion. Uploading, on its own, makes no
    LLM call and consumes NOTHING of the caller's metered run allowance.

    U2a (R-F1/R-F3/MON-012) — what the upload now does with the FILE, not just
    its text:

    * The exact bytes are stored on the new résumé (``originalFile``) with
      their name and their VERIFIED content type, and ``formatHash`` becomes the
      full SHA-256 of those bytes. That upload is the user's immutable baseline
      document: it is written once here and no code path ever rewrites it, so
      every later tailoring run derives from a source that still exists.
      ``GET /resumes/{id}/original`` serves it back byte-identical.
    * Format is decided by CONTENT, not by the client's claims: ``%PDF`` magic
      → PyMuPDF, an OOXML ZIP containing ``word/`` → python-docx paragraph and
      table text, ``.txt``/``.md`` → strict UTF-8. Anything else is a 422 that
      names the supported formats. Previously every non-PDF upload was decoded
      with ``errors="replace"``, so a real .docx, a screenshot, or a corrupt
      file became a Resume row of U+FFFD garbage presented back as the user's
      résumé (MON-012).
    * Uploads over ``MAX_UPLOAD_BYTES`` are refused with 413 on a bounded read.

    ``extract_stories`` (F-03, PROD-UAT-2026-08-03) — OPT-IN, default OFF.
    This endpoint used to dispatch the ``storyExtractor`` agent
    unconditionally so the Story Bank would reflect the new base resume
    (SC-SB-01). That agent is genuine LLM work (STRUCTURED tier, one
    ``complete_json`` call per four résumé bullets), so it is metered: a
    single deliberate upload silently produced an unrequested agent run
    (``costUsd 0.0010``, ``billingAudit.quotaPath "metered_api"``) and burned
    one of a Free plan's five monthly runs, with no warning before the fact
    and no way to decline. Exempting the agent from metering was the wrong
    remedy — it really does call a model, and the exemption seam
    (``_DETERMINISTIC_BACKENDS`` / ``_OPTIONAL_LLM_BY_BACKEND``) exists only
    for calls that reach NO model — so extraction is now the caller's explicit
    choice, priced and disclosed before they commit. When it IS requested the
    dispatch is unchanged: same agent, same atomic reserve, same audit row,
    same GAP-P6-RESFIX entitlement propagation. The capability itself is not
    lost — ``POST /agents/story-extractor/run`` (the Story Bank's "Draft
    missing stories" trigger) runs exactly the same extraction on demand.

    The response reports both halves honestly: ``storyExtractionRequested``
    says whether extraction was asked for, and ``storyExtraction`` carries the
    run result — ``None`` when it was not requested, so no caller can render
    after-the-fact copy claiming a run that never happened.
    """
    # Bounded read: one byte past the cap is enough to prove the file is over
    # it, so an oversized upload is never fully buffered or persisted.
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            f"Resume file is larger than the {MAX_UPLOAD_BYTES // (1024 * 1024)}MB "
            "upload limit.",
        )
    filename = _safe_upload_filename(file.filename or "resume")
    raw_text, content_type = _extract_upload_text(filename, data)
    if len(raw_text.strip()) < 50:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Extracted resume text is too short to be a resume",
        )
    from app.services.resume_tailor import extract_bullets

    stem = filename.rsplit(".", 1)[0][:100] or "Uploaded resume"
    sections = {
        "raw_text": raw_text,
        "bullets": [
            {"text": b, "evidenceRef": f"bullet-{i}"}
            for i, b in enumerate(extract_bullets(raw_text))
        ],
        "contact": {},
    }
    repo = ResumeRepository()
    resume = repo.create(
        current_user["id"],
        sections,
        # FULL SHA-256 of the stored bytes (U2a / R-F1). This used to be
        # truncated to 16 hex chars; the digest now identifies a document we
        # actually keep, and matches resume_parser.compute_format_hash's
        # long-standing convention. resolve_original_pdf accepts both widths,
        # so pre-existing truncated hashes keep resolving.
        hashlib.sha256(data).hexdigest(),
        label=f"Uploaded — {stem}",
        version=repo.next_version(current_user["id"]),
        original_file=data,
        original_filename=filename,
        original_content_type=content_type,
    )
    extraction: dict[str, Any] | None = None
    if extract_stories:
        try:
            from app.routers.agents import _dispatch

            extraction = _dispatch(current_user["id"], "storyExtractor", {})
        except HTTPException:
            # An HTTPException here (e.g. the 402 subscription-required paywall
            # gate in _record_run) is a real API error, not an extraction
            # failure — it must propagate to the client so a non-subscriber is
            # routed to /pricing instead of getting a 200 with the error buried
            # in storyExtraction.error (GAP-P6-RESFIX).
            raise
        except Exception as exc:  # noqa: BLE001 — upload must survive extraction issues
            extraction = {"error": str(exc)}
    return {
        **resume,
        "storyExtraction": extraction,
        "storyExtractionRequested": extract_stories,
    }


@router.get("/{resume_id}")
def get_resume(resume_id: str, current_user: CurrentUser) -> dict[str, Any]:
    resume = ResumeRepository().get_by_id(resume_id, current_user["id"])
    if resume is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")
    return resume


@router.get("/{resume_id}/ats")
def ats_score(
    resume_id: str, current_user: CurrentUser, job_id: str | None = None
) -> dict[str, Any]:
    """Deterministic ATS score of this resume version against a job description.

    Scores against the version's source job by default (``?job_id=`` overrides).
    The breakdown is the real ATS engine output — keyword coverage, semantic
    similarity, experience gap — never a fabricated number (SC-RS-05).
    """
    repo = ResumeRepository()
    resume = repo.get_by_id(resume_id, current_user["id"])
    if resume is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")
    target_job_id = job_id or resume.get("sourceJobId")
    if not target_job_id:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "Resume has no target job — tailor it against a job or pass ?job_id=",
        )
    from app.repositories.job import JobRepository

    job = JobRepository().get_by_id(target_job_id, current_user["id"])
    if job is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Target job not found")
    sections = resume.get("sections") or {}
    text = sections.get("raw_text") or "\n".join(
        b.get("text", "") for b in sections.get("bullets", [])
    )
    if not text.strip():
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY, "Resume has no scoreable text"
        )
    from app.services.ats_engine import ATSEngine
    from app.services.fit_evidence import job_evidence_text

    # F-UAX-01: score against the SAME JD text `GET /jobs/{id}/insights` uses
    # (title + description + requirements, `job_evidence_text`) — not the
    # description alone. Both endpoints feed the Resume Studio before/after
    # panel for the same job, so a JD-corpus mismatch here would leak into
    # every displayed ATS and dimension delta as a spurious term that has
    # nothing to do with the tailoring itself.
    score = ATSEngine().score(text, job_evidence_text(job))
    return {
        "resume_id": resume_id,
        "job_id": target_job_id,
        "job_title": job.get("title"),
        "company": job.get("company"),
        "overall": round(score.overall, 1),
        "keyword_match": round(score.keyword_match, 1),
        "semantic_similarity": round(score.semantic_similarity, 1),
        # GMV4-ats-002: which path actually produced semantic_similarity —
        # "local"/"hf_api" (genuine) or "degraded" (neutral placeholder, see
        # ats_engine.py's ATSScore.semantic_path). Without this, a degraded
        # score is indistinguishable from a real measurement to any caller.
        "semantic_path": score.semantic_path,
        # Unambiguous, client-branchable boolean twin of semantic_path —
        # never requires the client to know the sentinel string. Round 3:
        # WHITELIST — only "local"/"hf_api" count as measured; "degraded",
        # "untracked" and any unrecognised value all read as degraded.
        "semantic_degraded": score.semantic_path not in ("local", "hf_api"),
        "experience_gap": round(score.experience_gap, 1),
        "matched_keywords": score.matched_keywords,
        "missing_keywords": score.missing_keywords,
        "requires_review": score.requires_review,
    }


def _resume_scoreable_text(resume: dict[str, Any]) -> str:
    sections = resume.get("sections") or {}
    return sections.get("raw_text") or "\n".join(
        b.get("text", "") for b in sections.get("bullets", [])
    )


@router.get("/{resume_id}/tailoring-impact")
def tailoring_impact(
    resume_id: str, current_user: CurrentUser, job_id: str | None = None
) -> dict[str, Any]:
    """The before/after pair for one tailored version — ONE authority (R-01/R-03).

    U-AX build spec item 3 ("BEFORE/AFTER HONESTY"). Both halves are produced
    HERE, by the same blend and the same rounding authority
    (``routers/jobs.py::build_fit_dimensions`` / ``_round``: integer, clamped
    [0,100]), against the same JD corpus (``job_evidence_text``):

    * ``before`` — the baseline résumé vs this job, delegated verbatim to
      ``_build_insights``, i.e. the identical payload the Job Discovery fit
      radar renders. Not a re-implementation; literally the same call.
    * ``after`` — the tailored version's own re-score, blended by the same
      function.

    Round 2 computed the "after" half in the BROWSER from the wire's
    already-1-decimal subscores while the "before" half arrived pre-rounded to
    integers server-side. That duplicated blend put up to ±0.5 of rounding
    artefact into every displayed delta (~25% of the product's measured ~2-point
    average lift) and had to re-derive the provenance rules by hand, which is
    how a placeholder-contaminated baseline reached the screen flagged as
    measured. There is now nothing left to duplicate.

    PROVENANCE: a half whose ATS is not a genuine measurement reports
    ``ats: null`` + ``atsMeasured: false``. The number is WITHHELD rather than
    flagged, so no consumer can render it as a bold headline the way the panel
    did in round 2 — the untrustworthy arm simply carries no number.
    """
    from app.repositories.job import JobRepository
    from app.routers.jobs import _build_insights, _round, build_fit_dimensions
    from app.services.ats_engine import ATSEngine
    from app.services.fit_evidence import job_evidence_text

    repo = ResumeRepository()
    resume = repo.get_by_id(resume_id, current_user["id"])
    if resume is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")
    target_job_id = job_id or resume.get("sourceJobId")
    if not target_job_id:
        # Picking a job would fabricate the comparison: a before/after pair is
        # only meaningful against the posting the version was tailored for.
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Resume has no target job — tailor it against a job or pass ?job_id=",
        )
    job = JobRepository().get_by_id(target_job_id, current_user["id"])
    if job is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Target job not found")
    text = _resume_scoreable_text(resume)
    if not text.strip():
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY, "Resume has no scoreable text"
        )

    before = _build_insights(job, current_user["id"])
    before_measured = bool(
        before.get("scored")
        and before.get("atsMeasured")
        and not before.get("semanticDegraded")
    )

    try:
        score = ATSEngine().score(text, job_evidence_text(job))
        after_trusted = score.semantic_path in ("local", "hf_api")
        after_dimensions = build_fit_dimensions(
            job,
            keyword_match=float(score.keyword_match),
            semantic=float(score.semantic_similarity),
            experience=float(score.experience_gap),
            overall=float(score.overall),
            semantic_trusted=after_trusted,
            resume_measured=True,
        )
        after_ats = _round(float(score.overall)) if after_trusted else None
        after_measured = after_trusted
        after_reason = (
            None
            if after_trusted
            else "the semantic scoring model was unavailable for this run"
        )
    except Exception:  # noqa: BLE001 — an unscoreable half is reported, never faked
        # No engine output at all: unlike the job panel there is no stored
        # fitScore to fall back on for a tailored version, and inventing one
        # is exactly the defect this endpoint exists to close.
        after_dimensions = build_fit_dimensions(
            job,
            keyword_match=0.0, semantic=0.0, experience=0.0, overall=0.0,
            semantic_trusted=False, resume_measured=False,
        )
        after_ats = None
        after_measured = False
        after_reason = "the ATS scoring engine could not score this version"

    before_reason = (
        None
        if before_measured
        else (
            "the ATS scoring engine could not score your baseline résumé "
            "against this posting"
            if not before.get("atsMeasured")
            else "the semantic scoring model was unavailable for the baseline score"
        )
    )
    return {
        "resumeId": resume_id,
        "jobId": target_job_id,
        "jobTitle": job.get("title"),
        "company": job.get("company"),
        # Stated on the wire so a client cannot reintroduce a second rounding
        # hop without contradicting the payload it is rendering.
        "granularity": "integer_0_100",
        "before": {
            "label": "Baseline résumé",
            "ats": before["overall"] if before_measured else None,
            "atsMeasured": before_measured,
            "unmeasuredReason": before_reason,
            "dimensions": before["dimensions"],
        },
        "after": {
            "label": "Tailored version",
            "ats": after_ats,
            "atsMeasured": after_measured,
            "unmeasuredReason": after_reason,
            "dimensions": after_dimensions,
        },
    }


@router.get("/{resume_id}/diff")
def diff_resume(resume_id: str, current_user: CurrentUser) -> dict[str, Any]:
    """Bullet-level diff of a tailored resume against its parent."""
    repo = ResumeRepository()
    resume = repo.get_by_id(resume_id, current_user["id"])
    if resume is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")
    if not resume.get("parentId"):
        return {"resume_id": resume_id, "parent_id": None, "changes": []}
    parent = repo.get_by_id(resume["parentId"], current_user["id"])
    # First occurrence wins for duplicated refs (pre-fix tailored versions),
    # matching the tailor service's healing — the diff stays consistent with
    # the originals each rewrite was actually validated against.
    parent_by_ref: dict[Any, str] = {}
    for b in (parent or {}).get("sections", {}).get("bullets", []):
        parent_by_ref.setdefault(b.get("evidenceRef"), b.get("text", ""))
    changes = []
    for bullet in resume.get("sections", {}).get("bullets", []):
        ref = bullet.get("evidenceRef")
        original = parent_by_ref.get(ref, "")
        if bullet.get("text") != original:
            changes.append(
                {"evidenceRef": ref, "before": original, "after": bullet.get("text")}
            )
    return {"resume_id": resume_id, "parent_id": resume["parentId"], "changes": changes}


def _branded_content(
    resume: dict[str, Any],
) -> tuple[str, str, str, list[dict[str, Any]]]:
    """Map a stored resume record onto the branded template's inputs.

    Used only on the structured-rendering fallback (no bundled source PDF on
    disk), so the resume is rebuilt from its ``sections`` payload rather than
    edited in place. Name/title/objective come from the parsed contact block
    when present, with the resume label and the first content line as
    fallbacks; every bullet is grouped under a single Experience heading.
    """
    sections = resume.get("sections", {}) or {}
    contact = sections.get("contact", {}) or {}
    raw_lines = [
        line.strip() for line in str(sections.get("raw_text", "")).splitlines() if line.strip()
    ]
    name = str(
        contact.get("name")
        or (raw_lines[0] if raw_lines else "")
        or resume.get("label")
        or "Resume"
    )
    title = str(contact.get("title") or contact.get("headline") or "")
    objective = str(sections.get("objective") or sections.get("summary") or "")
    bullets = [
        str(b.get("text", ""))
        for b in sections.get("bullets", [])
        if str(b.get("text", "")).strip()
    ]
    template_sections = [{"heading": "Experience", "bullets": bullets}] if bullets else []
    return name, title, objective, template_sections


@router.get("/{resume_id}/download")
def download_resume(resume_id: str, current_user: CurrentUser) -> Response:
    """Download a resume as a format-preserving PDF.

    - **Base resume** (no parent) backed by a bundled asset: the original
      bundled PDF bytes, verbatim.
    - **Tailored resume**: the original PDF with *only* the reworded bullets
      redrawn in place — same two-column layout, peach title panel, coral
      accents and fonts — plus a subtle highlight behind each changed bullet.
      Every unchanged element stays byte-for-byte identical to the source.
    - **No bundled source PDF on disk** (e.g. an externally-ingested variant):
      the resume is rebuilt from its structured content with the branded
      two-page template — each reworded bullet washed coral on page 2.
    """
    from app.services.resume_pdf import (
        create_branded_resume_pdf,
        render_tailored_pdf,
        resolve_original_pdf,
    )

    repo = ResumeRepository()
    resume = repo.get_by_id(resume_id, current_user["id"])
    if resume is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")

    parent_id = resume.get("parentId")
    parent = repo.get_by_id(parent_id, current_user["id"]) if parent_id else None

    # The reworded bullets, diffed against the parent (empty for a base resume).
    changes: list[tuple[str, str]] = []
    if parent is not None:
        parent_by_ref = {
            b.get("evidenceRef"): b.get("text", "")
            for b in parent.get("sections", {}).get("bullets", [])
        }
        for bullet in resume.get("sections", {}).get("bullets", []):
            before = parent_by_ref.get(bullet.get("evidenceRef"))
            after = bullet.get("text", "")
            if before and after and before != after:
                changes.append((before, after))

    original = resolve_original_pdf(
        (parent or resume).get("formatHash") or resume.get("formatHash")
    )

    if original is not None and original.exists():
        # A bundled source PDF backs THIS résumé (seeded base / BA variant) →
        # preserve its exact layout.
        if parent is None:
            pdf_bytes = original.read_bytes()  # base → verbatim bytes
        else:
            pdf_bytes = render_tailored_pdf(original, changes)  # splice in place
    else:
        # No bundled source PDF for this résumé — a user-authored upload/ingest
        # (NF-final-B-005) → render from the résumé's OWN structured content with
        # the branded template; never serve the operator's bundled PDF bytes.
        name, title, objective, sections = _branded_content(resume)
        pdf_bytes = create_branded_resume_pdf(
            name, title, objective, sections, changes or None
        )

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="resume-{resume_id[:8]}.pdf"'},
    )


@router.get("/{resume_id}/original")
def download_original_resume(resume_id: str, current_user: CurrentUser) -> Response:
    """Download the résumé EXACTLY as it was uploaded (U2a / R-F1).

    This is the immutable baseline document — the same bytes, the same content
    type, the same file name the user gave us. It is deliberately distinct from
    ``/download``, which RENDERS a résumé (verbatim bundled PDF, in-place
    tailored splice, or the branded template) and can therefore return a
    document that never existed on the user's disk.

    Owner-only, exactly like every other ``/resumes/{id}`` route: another
    account's résumé is indistinguishable from a non-existent one (404).

    A résumé with no stored bytes 404s with an honest explanation instead of
    synthesising a stand-in file. That is the state of every row created before
    this slice (original bytes were never kept), of every JSON-ingested résumé
    (``POST /resumes`` carries text, not a file), and of every tailored child
    version (produced by the pipeline, never uploaded).
    """
    record = ResumeRepository().get_original_file(resume_id, current_user["id"])
    if record is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Resume not found")
    data = record["originalFile"]
    if not data:
        raise HTTPException(
            status.HTTP_404_NOT_FOUND,
            "No original file is stored for this résumé. It was created without "
            "a file upload (typed/ingested text or a tailored version), or it "
            "was uploaded before Aether began preserving original documents — "
            "those bytes no longer exist, so there is nothing to return.",
        )
    filename = _safe_upload_filename(record["originalFilename"] or f"resume-{resume_id[:8]}")
    return Response(
        content=data,
        media_type=record["originalContentType"] or "application/octet-stream",
        headers={"Content-Disposition": _content_disposition(filename)},
    )