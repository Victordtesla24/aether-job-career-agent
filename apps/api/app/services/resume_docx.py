"""Format-preserving DOCX tailoring engine (U2b / U-PLAN.md R-F4).

The baseline résumé the user uploads in Settings is the immutable source of
truth for CONTENT **and FORMAT**. Tailoring may therefore only change words —
never structure, styles, fonts, spacing, colours or any other visual element.

For a Word document that is achievable exactly, because a ``.docx`` is a
structured OOXML package rather than a page image: this module opens the
user's own stored bytes and rewrites **only the text of the runs that make up
a changed bullet**, leaving every other part of the package — ``styles.xml``,
``numbering.xml``, section/page setup, headers, tables, images, and every
untouched paragraph's own runs — exactly as the user authored it. That is the
"native run-level text replacement" R-F4 names as the flagship path, and it is
the opposite of the pre-U2b behaviour, where every real upload was re-flowed
into Aether's generic branded PDF template while the UI claimed preservation
(MON-011).

Two properties are load-bearing and are asserted by the tests:

* **The baseline is never mutated.** ``render_tailored_docx`` reads a COPY of
  the caller's buffer and returns new bytes; the stored original is
  write-once at upload (see ``repositories/resume.py``) and no code path here
  writes back to it.
* **Run formatting survives a rewrite.** A bullet whose first words are bold
  and whose remainder is regular text keeps that structure: the replacement
  text is distributed back across the paragraph's existing runs by diffing old
  against new (``difflib``) and mapping each run's character span through the
  diff, instead of the naive "dump everything into run 0" approach that would
  bold an entire line.

ATS-safety (product rule): the output stays a real text layer — no image, text
box or OLE substitute is ever introduced — so an ATS parser reads the tailored
words exactly as it read the baseline's.
"""
from __future__ import annotations

import zipfile
from difflib import SequenceMatcher
from io import BytesIO
from typing import Any, Callable, Iterator, Sequence

#: Local-file-header / empty-archive / spanned-archive ZIP signatures. A .docx
#: is an OOXML ZIP, so this is the cheap first gate before opening it.
_ZIP_MAGICS = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

#: Leading list-marker glyphs (and the whitespace around them) that a résumé
#: bullet may carry as literal text. Stripped only from the START of a line, so
#: a hyphen inside the sentence ("test-evidence") is never touched.
_MARKER_PREFIX = "•●▪◦‣·-–—*•●▪ \t"

#: The marker Aether's own text model uses for a list item. Word usually stores
#: bullets as numbering properties rather than characters, so the glyph does not
#: appear in ``paragraph.text`` at all; emitting it during extraction is what
#: lets ``resume_tailor.extract_bullets`` (a line-marker state machine) see a
#: Word bullet as a bullet at all.
_LIST_MARKER = "• "


class DocxParseError(RuntimeError):
    """A .docx package that cannot be read as a WordprocessingML document."""


def looks_like_docx(data: bytes) -> bool:
    """True when ``data`` really is an OOXML word-processing package.

    Checked by CONTENT, never by a filename or a client-supplied Content-Type:
    the bytes must be a readable ZIP containing the WordprocessingML part.
    """
    if not data or not data.startswith(_ZIP_MAGICS):
        return False
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            return any(name.startswith("word/") for name in archive.namelist())
    except (zipfile.BadZipFile, OSError):
        return False


def _open(data: bytes) -> Any:
    """Open ``data`` as a python-docx ``Document`` from an independent copy.

    ``bytes(data)`` guarantees the caller's buffer is never handed to (or
    mutated by) the parser even when it passes a ``bytearray``/``memoryview``
    straight out of the database driver.
    """
    from docx import Document

    return Document(BytesIO(bytes(data)))


def _iter_paragraphs(container: Any) -> Iterator[Any]:
    """Every paragraph in the document body, including nested table cells.

    Résumés are frequently laid out in tables (the two-column style is almost
    always a table), so a body-only walk would silently skip most bullets and
    leave the tailored download identical to the baseline.
    """
    for paragraph in container.paragraphs:
        yield paragraph
    for table in getattr(container, "tables", ()):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _run_text(paragraph: Any) -> str:
    """The paragraph's text as the RUNS spell it.

    Deliberately not ``paragraph.text``: the runs are what this module rewrites,
    so matching and splicing must agree on the same string.
    """
    return "".join(run.text for run in paragraph.runs)


def _is_list_paragraph(paragraph: Any) -> bool:
    """True when Word itself treats this paragraph as a list item.

    Either through a list style ("List Bullet", "List Paragraph", "List
    Number") or through direct numbering properties (``w:numPr``) — the two
    ways a .docx encodes a bullet without putting a glyph in the text.
    """
    style_name = ""
    try:
        style_name = (paragraph.style.name or "").lower()
    except (AttributeError, KeyError):  # pragma: no cover - malformed style ref
        style_name = ""
    if "list" in style_name:
        return True
    p_pr = getattr(paragraph._p, "pPr", None)
    return p_pr is not None and getattr(p_pr, "numPr", None) is not None


def extract_docx_lines(data: bytes) -> list[str]:
    """Résumé text lines from a .docx, with Word's own list items marked.

    Paragraph order and blank lines are preserved because
    :func:`app.services.resume_tailor.extract_bullets` reads line structure
    (marker lines, all-caps section banners) to reassemble bullets. A paragraph
    Word treats as a list item is prefixed with ``"• "`` so it survives that
    state machine: Word stores the bullet glyph in ``numbering.xml``, not in the
    paragraph text, so before this every DOCX résumé extracted ZERO bullets and
    was therefore untailorable — the upload succeeded and then nothing could be
    rewritten.

    Table cells are appended after the body paragraphs (``document.paragraphs``
    covers only body-level paragraphs, so nothing is emitted twice).
    """
    document = _open(data)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text and _is_list_paragraph(paragraph):
            text = f"{_LIST_MARKER}{text}"
        lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return lines


def _normalize(text: str) -> str:
    """Collapse all whitespace so a soft line break can't defeat a match."""
    return " ".join(text.split())


def _strip_marker(text: str) -> str:
    """The bullet's own sentence, without a leading list glyph."""
    return text.lstrip(_MARKER_PREFIX).strip()


def _marker_prefix(text: str) -> str:
    """The literal leading marker/indent of ``text`` (possibly empty)."""
    stripped = text.lstrip(_MARKER_PREFIX)
    return text[: len(text) - len(stripped)]


def _offset_mapper(old: str, new: str) -> Callable[[int], int]:
    """Map a character offset in ``old`` onto the matching offset in ``new``.

    Monotonic by construction: inside an equal block the offset shifts by the
    block's own displacement; inside a replaced/deleted block every offset
    snaps to the start of its replacement, so an edit is attributed wholly to
    the run that contained it and no character of ``new`` is ever dropped or
    duplicated when consecutive run spans are sliced with it.
    """
    opcodes = SequenceMatcher(None, old, new, autojunk=False).get_opcodes()

    def mapper(pos: int) -> int:
        if pos <= 0:
            return 0
        if pos >= len(old):
            return len(new)
        for tag, i1, i2, j1, _j2 in opcodes:
            if i1 <= pos < i2:
                return j1 + (pos - i1) if tag == "equal" else j1
            if pos < i1:
                return j1
        return len(new)

    return mapper


def _rewrite_paragraph(paragraph: Any, new_text: str) -> bool:
    """Replace the paragraph's text, keeping its runs and their formatting.

    Returns ``True`` when something actually changed. Every run object survives
    (so bold lead-ins, colours, fonts, highlights and character styles survive
    with it); only the characters inside them move.
    """
    runs = paragraph.runs
    if not runs:
        # Nothing to write into without INVENTING a run — and an invented run
        # would carry default formatting, which is exactly the silent
        # re-formatting this engine exists to prevent.
        return False
    old_text = "".join(run.text for run in runs)
    if old_text == new_text:
        return False
    mapper = _offset_mapper(old_text, new_text)
    cursor = 0
    for run in runs:
        start, end = cursor, cursor + len(run.text)
        cursor = end
        run.text = new_text[mapper(start) : mapper(end)]
    return True


def apply_docx_changes(
    document: Any, changes: Sequence[tuple[str, str]]
) -> list[tuple[str, str]]:
    """Apply ``(before, after)`` bullet rewrites to ``document`` in place.

    Matching is deliberately conservative — a change that cannot be located
    with certainty is skipped and reported as unapplied rather than guessed at,
    because a wrong splice silently corrupts the user's own document:

    1. **Whole-paragraph match** on the normalised, marker-stripped text. This
       is the ordinary case: a résumé bullet is one paragraph.
    2. **Literal containment** inside a paragraph, for the case where a bullet
       shares its paragraph with other text.

    Each paragraph is consumed at most once, so two bullets with identical text
    map to two different paragraphs instead of both rewriting the first.

    Returns the changes that were genuinely applied, in input order — the
    caller needs that to report fidelity HONESTLY instead of assuming success.
    """
    paragraphs = list(_iter_paragraphs(document))
    used: set[int] = set()
    applied: list[tuple[str, str]] = []
    pending: list[tuple[str, str]] = []

    for before, after in changes:
        before_core = _strip_marker(_normalize(before))
        after_core = _strip_marker(_normalize(after))
        if not before_core or not after_core or before_core == after_core:
            continue
        matched = False
        for index, paragraph in enumerate(paragraphs):
            if index in used:
                continue
            raw = _run_text(paragraph)
            if _strip_marker(_normalize(raw)) != before_core:
                continue
            if _rewrite_paragraph(paragraph, f"{_marker_prefix(raw)}{after_core}"):
                used.add(index)
                applied.append((before, after))
            matched = True
            break
        if not matched:
            pending.append((before, after))

    for before, after in pending:
        before_core = _strip_marker(_normalize(before))
        after_core = _strip_marker(_normalize(after))
        for index, paragraph in enumerate(paragraphs):
            if index in used:
                continue
            raw = _run_text(paragraph)
            if before_core not in raw:
                continue
            if _rewrite_paragraph(paragraph, raw.replace(before_core, after_core, 1)):
                used.add(index)
                applied.append((before, after))
            break
    return applied


def render_tailored_docx_report(
    original: bytes, changes: Sequence[tuple[str, str]]
) -> tuple[bytes, list[tuple[str, str]]]:
    """``(tailored_bytes, applied_changes)`` — the honest form of the render.

    Callers that must tell the user what really happened (the download
    endpoint's fidelity report) use this; :func:`render_tailored_docx` is the
    same operation for callers that only need the document.
    """
    try:
        document = _open(original)
    except Exception as exc:  # noqa: BLE001 — every layer raises its own type
        raise DocxParseError(str(exc)) from exc
    applied = apply_docx_changes(document, changes)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), applied


def render_tailored_docx(original: bytes, changes: Sequence[tuple[str, str]]) -> bytes:
    """The user's own .docx with ONLY the reworded bullets rewritten (R-F4).

    ``original`` is never mutated — the baseline document stays byte-identical
    in storage and in memory. Every element the rewrite does not name (styles,
    numbering, section setup, headers/footers, tables, images, and every other
    paragraph's runs) is carried through untouched by python-docx's own package
    round-trip.
    """
    rendered, _applied = render_tailored_docx_report(original, changes)
    return rendered
